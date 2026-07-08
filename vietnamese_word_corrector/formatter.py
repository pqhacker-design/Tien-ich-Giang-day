import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def set_paragraph_text_and_style(p, text, font_size=13, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, uppercase=False):
    """Hàm bổ trợ: Xóa bỏ các Run cũ bị lỗi style và tái tạo định dạng chuẩn 100%"""
    if uppercase:
        text = text.upper()
        
    # Xóa các run cũ bị xé nhỏ
    p.text = "" 
    p.alignment = align
    
    # Tạo lại 1 run duy nhất chuẩn Times New Roman
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0) # Màu đen chuẩn

def normalize_to_nd30(input_docx_path, output_docx_path):
    """
    Tự động chuẩn hóa lề lối, font chữ, thụt lề, giãn dòng, IN HOA và TÔ ĐẬM 
    các đề mục theo quy định Nghị định 30/2020/NĐ-CP.
    """
    doc = docx.Document(input_docx_path)

    # 1. Định dạng Lề trang giấy theo NĐ 30/2020/NĐ-CP
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)

    # 2. Các quy tắc Biểu thức chính quy (Regex) mở rộng
    pattern_quoc_hieu = re.compile(r'^(CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM|ĐỘC LẬP - TỰ DO - HẠNH PHÚC)', re.IGNORECASE)
    pattern_main_title = re.compile(r'^(BẢNG THUYẾT MINH|KẾT LUẬN|TỔNG QUAN|QUYẾT ĐỊNH|THÔNG BÁO|BÁO CÁO|TỜ TRÌNH|KẾ HOẠCH)', re.IGNORECASE)
    
    # Bắt các đề mục la mã (I., II., 1., 2.)
    pattern_roman_heading = re.compile(r'^(I|II|III|IV|V|VI|VII|VIII|IX|X)[\.\s\:]') 
    pattern_arabic_heading = re.compile(r'^\d+[\.\s\)]')                           
    pattern_sub_heading = re.compile(r'^[a-z]\)')                                 
    
    # Bắt các đề mục từ khóa đặc thù (Ý nghĩa:, Tông màu..., Phân tích..., Sự kiện:...)
    pattern_key_heading = re.compile(r'^(SỰ KIỆN|Ý NGHĨA|PHÂN TÍCH|KẾT LUẬN|TÔNG MÀU|PHÔNG CHỮ|SỰ KẾT HỢP)[\:\s]', re.IGNORECASE)

    for p in doc.paragraphs:
        text_raw = p.text.strip()
        if not text_raw:
            continue

        # Cấu hình giãn dòng chung
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

        # --- A. QUỐC HIỆU / TIÊU ĐỀ LỚN CHÍNH (IN HOA, IN ĐẬM, CĂN GIỮA) ---
        if pattern_quoc_hieu.search(text_raw) or (doc.paragraphs.index(p) == 0 and len(text_raw) < 80):
            p.paragraph_format.first_line_indent = Cm(0)
            set_paragraph_text_and_style(p, text_raw, font_size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, uppercase=True)
            continue

        if pattern_main_title.match(text_raw):
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(6)
            set_paragraph_text_and_style(p, text_raw, font_size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, uppercase=True)
            continue

        # --- B. ĐỀ MỤC LA MÃ & ĐỀ MỤC ĐÁNH SỐ (1., 2., I., II.) -> IN HOA & IN ĐẬM ---
        if pattern_roman_heading.match(text_raw) or pattern_arabic_heading.match(text_raw):
            p.paragraph_format.first_line_indent = Cm(0) # Đề mục không thụt lề
            p.paragraph_format.space_before = Pt(6)
            set_paragraph_text_and_style(p, text_raw, font_size=13, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, uppercase=True)
            continue

        # --- C. CÁC ĐỀ MỤC TỪ KHÓA (Sự kiện:, Ý nghĩa:, Tông màu...) -> IN ĐẬM ---
        if pattern_key_heading.match(text_raw) or pattern_sub_heading.match(text_raw):
            p.paragraph_format.first_line_indent = Cm(0.5)
            set_paragraph_text_and_style(p, text_raw, font_size=13, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, uppercase=False)
            continue

        # --- D. VĂN BẢN THƯỜNG (BODY TEXT) -> CĂN ĐỀU, THỤT ĐẦU DÒNG 1.27 CM ---
        p.paragraph_format.first_line_indent = Cm(1.27) # Thụt đầu dòng chuẩn 1.27cm
        set_paragraph_text_and_style(p, text_raw, font_size=13, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, uppercase=False)

    # 3. Định dạng lại cả Bảng (Tables) nếu có trong file Word
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    if p.text.strip():
                        set_paragraph_text_and_style(p, p.text.strip(), font_size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.save(output_docx_path)
    return output_docx_path
