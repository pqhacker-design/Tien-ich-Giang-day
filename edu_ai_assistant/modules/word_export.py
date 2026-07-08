import io
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class WordExportEngine:
    @staticmethod
    def create_docx(text_content: str) -> io.BytesIO:
        doc = docx.Document()
        
        # 1. Cấu hình lề chuẩn Nghị định 30/2020/NĐ-CP (Trên 2cm, Dưới 2cm, Trái 3cm, Phải 2cm)
        for section in doc.sections:
            section.top_margin = Inches(0.79)      # 2 cm
            section.bottom_margin = Inches(0.79)   # 2 cm
            section.left_margin = Inches(1.18)     # 3 cm
            section.right_margin = Inches(0.79)    # 2 cm

        # 2. Thiết lập Style mặc định (Times New Roman, 13pt)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(13)
        font.color.rgb = RGBColor(0, 0, 0)

        lines = text_content.split("\n")
        in_table = False
        table_rows = []

        for line in lines:
            raw_line = line.strip()
            if not raw_line or raw_line.startswith("---"):
                continue

            # Xử lý bảng (Table) nếu có cấu trúc | ... |
            if raw_line.startswith("|") and raw_line.endswith("|"):
                # Bỏ qua dòng phân cách bảng (ví dụ: | :--- | :--- |)
                if ":---" in raw_line or "---" in raw_line:
                    continue
                cells = [c.strip() for c in raw_line.split("|")[1:-1]]
                table_rows.append(cells)
                in_table = True
                continue
            else:
                # Nếu kết thúc bảng, vẽ bảng vào Word
                if in_table and table_rows:
                    cols_count = max(len(r) for r in table_rows)
                    tbl = doc.add_table(rows=len(table_rows), cols=cols_count)
                    tbl.autofit = True
                    for r_idx, row in enumerate(table_rows):
                        for c_idx, cell_text in enumerate(row):
                            if c_idx < cols_count:
                                cell = tbl.cell(r_idx, c_idx)
                                # Lọc sạch Markdown in đậm trong bảng
                                clean_cell_text = cell_text.replace("**", "").replace("*", "")
                                p = cell.paragraphs[0]
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = p.add_run(clean_cell_text)
                                if "**" in cell_text or r_idx == 0:
                                    run.bold = True
                    table_rows = []
                    in_table = False

            # Xử lý danh sách có dấu gạch đầu dòng (Bullet points)
            is_bullet = False
            clean_line = raw_line
            if raw_line.startswith("* ") or raw_line.startswith("* "):
                is_bullet = True
                clean_line = raw_line.lstrip("* ").strip()

            # Tạo đoạn văn bản
            if is_bullet:
                p = doc.add_paragraph(style='List Bullet')
            else:
                p = doc.add_paragraph()

            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(3)

            # Căn lề và định dạng theo tiêu đề
            if clean_line.isupper() and len(clean_line) < 100:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif clean_line.startswith("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM") or clean_line.startswith("Độc lập - Tự do"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif is_bullet:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Phân tích định dạng In đậm (**text**) trong câu
            parts = clean_line.split("**")
            for i, part in enumerate(parts):
                if not part:
                    continue
                # Loại bỏ các dấu * còn sót
                part_text = part.replace("*", "")
                run = p.add_run(part_text)
                # Các phần ở chỉ số lẻ giữa dấu ** sẽ được in đậm
                if i % 2 == 1:
                    run.bold = True

        # Render nốt bảng nếu bảng ở cuối file
        if in_table and table_rows:
            cols_count = max(len(r) for r in table_rows)
            tbl = doc.add_table(rows=len(table_rows), cols=cols_count)
            for r_idx, row in enumerate(table_rows):
                for c_idx, cell_text in enumerate(row):
                    if c_idx < cols_count:
                        cell = tbl.cell(r_idx, c_idx)
                        clean_cell_text = cell_text.replace("**", "").replace("*", "")
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(clean_cell_text)
                        if "**" in cell_text or r_idx == 0:
                            run.bold = True

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
