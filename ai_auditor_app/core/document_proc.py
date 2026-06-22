import docx
import pdfplumber
import io

class DocumentProcessor:
    @staticmethod
    def read_docx(file_bytes) -> str:
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)

    @staticmethod
    def read_pdf(file_bytes) -> str:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            text = "".join([page.extract_text() or "" for page in pdf.pages])
        return text

    @staticmethod
    def read_file(file) -> str:
        try:
            # Đã thêm: Đưa con trỏ file về vị trí ban đầu trước khi đọc
            file.seek(0) 
            
            if file.name.endswith('.docx'):
                return DocumentProcessor.read_docx(file.read())
            elif file.name.endswith('.pdf'):
                return DocumentProcessor.read_pdf(file.read())
            else:
                return "LỖI: Định dạng file không được hỗ trợ (Chỉ nhận .docx, .pdf)"
        except Exception as e:
            # Bọc lỗi ngăn chặn sập ứng dụng khi file hỏng hoặc lỗi định dạng zip
            return f"LỖI: Không thể đọc cấu trúc file '{file.name}'. Vui lòng kiểm tra lại định dạng file hoặc lưu lại file (Save as) bằng Microsoft Word chuẩn."

    @staticmethod
    def create_audited_docx(original_text: str, suggestions: list, dynamic_content: str = "") -> io.BytesIO:
        """
        Khởi tạo tài liệu kết quả giữ nguyên định dạng hoặc xuất báo cáo hoàn thiện chuyên nghiệp.
        """
        doc = docx.Document()
        doc.add_heading('BÁO CÁO HOÀN THIỆN VĂN BẢN - AI ASSISTANT', level=1)
        
        if dynamic_content:
            doc.add_heading('I. Nội Dung Khắc Phục Tự Động Từ AI', level=2)
            doc.add_paragraph(dynamic_content)
            
        doc.add_heading('II. Nhật Ký Đánh Giá Chi Tiết', level=2)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Tiêu chí / Yêu cầu'
        hdr_cells[1].text = 'Hiện trạng'
        hdr_cells[2].text = 'Đề xuất giải pháp'
        
        for sug in suggestions:
            row_cells = table.add_row().cells
            row_cells[0].text = str(sug.get('criteria', 'N/A'))
            row_cells[1].text = str(sug.get('status', 'N/A'))
            row_cells[2].text = str(sug.get('fix', 'N/A'))
            
        out_stream = io.BytesIO()
        doc.save(out_stream)
        out_stream.seek(0)
        return out_stream
