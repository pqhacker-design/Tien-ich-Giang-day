import io
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

class WordProcessor:
    @staticmethod
    def extract_text(file_bytes: bytes) -> str:
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = " ".join([p.text for p in cell.paragraphs if p.text.strip()])
                    if cell_text:
                        full_text.append(cell_text)
        return "\n".join(full_text)

    @staticmethod
    def insert_paragraph_after(paragraph, text, color_rgb, prefix=""):
        """Chèn đoạn văn mới liền sau paragraph chỉ định."""
        new_p = OxmlElement('w:p')
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        new_para.paragraph_format.space_before = Pt(2)
        new_para.paragraph_format.space_after = Pt(4)
        new_para.paragraph_format.line_spacing = 1.15
        run = new_para.add_run(f"{prefix} {text}" if prefix else text)
        run.font.color.rgb = color_rgb
        run.italic = True
        run.font.size = Pt(11)
        return new_para

    @staticmethod
    def integrate_digital_capacity(file_bytes: bytes, ai_data: dict, integration_type: str) -> io.BytesIO:
        doc = Document(io.BytesIO(file_bytes))
        sua_doi_list = ai_data.get('sua_doi', [])
        
        color_digital = RGBColor(0, 102, 204)   # Xanh dương
        color_ai = RGBColor(214, 107, 0)        # Vàng cam
        
        # Danh sách lưu các đoạn văn đã chèn để không chèn đè trùng lặp tại cùng 1 vị trí nếu không cần thiết
        used_paragraphs = set()

        for item in sua_doi_list:
            anchor = item.get('anchor_text', '').strip()
            content = item.get('insert_content', '').strip()
            loai = item.get('loai', 'Năng lực số')
            
            if not anchor or not content:
                continue
            
            prefix = "[Năng lực AI]:" if loai == "Năng lực AI" else "[Năng lực số]:"
            color = color_ai if loai == "Năng lực AI" else color_digital
            
            inserted = False
            
            # 1. Quét tìm trong danh sách Paragraphs
            for para in doc.paragraphs:
                # Điều kiện: Khớp anchor và đoạn này chưa bị chèn bởi thao tác trước đó
                if anchor in para.text and para not in used_paragraphs:
                    WordProcessor.insert_paragraph_after(para, content, color, prefix)
                    used_paragraphs.add(para)
                    inserted = True
                    break
            
            # 2. Nếu không có trong Paragraph thông thường, tìm tiếp trong Tables
            if not inserted:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if anchor in para.text and para not in used_paragraphs:
                                    WordProcessor.insert_paragraph_after(para, content, color, prefix)
                                    used_paragraphs.add(para)
                                    inserted = True
                                    break
                            if inserted:
                                break
                        if inserted:
                            break
                    if inserted:
                        break

        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        return output_stream
