import pandas as pd
from docx import Document
from io import BytesIO

class ExportEngine:
    @staticmethod
    def export_to_word(student_data: list) -> BytesIO:
        doc = Document()
        doc.add_heading('BÁO CÁO NHẬN XÉT HỌC SINH THCS 2026', 0)

        for item in student_data:
            doc.add_heading(f"Học sinh: {item.get('name')} - Lớp: {item.get('class_name')}", level=1)
            doc.add_paragraph(f"• Đánh giá Học tập: {item.get('academic_comment')}")
            doc.add_paragraph(f"• Đánh giá Phẩm chất: {item.get('quality_comment')}")
            doc.add_paragraph(f"• Đánh giá Năng lực: {item.get('capacity_comment')}")
            doc.add_paragraph(f"• Lời động viên / Hướng rèn luyện: {item.get('advice')}")
            doc.add_paragraph("-" * 40)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def export_to_excel(df: pd.DataFrame) -> BytesIO:
        buffer = BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='NhanXet_2026')
        buffer.seek(0)
        return buffer
