import docx
from pypdf import PdfReader
import io

def read_docx(file_bytes):
    """Đọc file DOCX và giữ cấu trúc văn bản cơ bản"""
    doc = docx.Document(io.BytesIO(file_bytes))
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Trích xuất văn bản trong bảng nếu có
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
                
    return "\n".join(full_text)

def read_pdf(file_bytes):
    """Đọc nội dung văn bản từ file PDF"""
    reader = PdfReader(io.BytesIO(file_bytes))
    full_text = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            full_text.append(text)
    return "\n".join(full_text)

def analyze_metadata(text):
    """AI phân tích sơ bộ Metadata hoặc dùng Regex thủ công nếu cần"""
    metadata = {
        "subject": "Chưa rõ",
        "grade": "Chưa rõ",
        "level": "THCS",
        "duration": "1 tiết",
        "topic": "Chưa rõ"
    }
    
    # Phân tích nhanh bằng heuristic từ dòng đầu tiên
    lines = [line.strip() for line in text.split('\n') if line.strip()][:10]
    combined_top = " ".join(lines).lower()
    
    if "toán" in combined_top: metadata["subject"] = "Toán học"
    elif "khoa học tự nhiên" in combined_top or "khtn" in combined_top: metadata["subject"] = "Khoa học tự nhiên"
    elif "vật lí" in combined_top: metadata["subject"] = "Vật lí"
    
    for g in range(1, 13):
        if f"lớp {g}" in combined_top or f"khối {g}" in combined_top:
            metadata["grade"] = f"Lớp {g}"
            if g <= 5: metadata["level"] = "Tiểu học"
            elif g <= 9: metadata["level"] = "THCS"
            else: metadata["level"] = "THPT"
            break
            
    return metadata