import docx
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import io

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Độ giãn biên trong của ô bảng để dữ liệu không bị dính sát viền"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def export_lesson_to_word(lesson_json, matrix_df, rubric_df, worksheet_data, lesson_format="Loại thường (Không chia cột)"):
    """Tạo file Microsoft Word .docx đạt chuẩn thể thức văn bản hành chính Việt Nam"""
    doc = docx.Document()
    
    # 1. Cấu hình Căn lề trang chuẩn (Top=2cm, Bottom=2cm, Left=3cm, Right=2cm)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        
        # Thiết lập Header & Footer tự động đánh số trang cơ bản
        footer = section.footer
        footer.paragraphs[0].text = "AI Giáo Án Thông Minh 4.0 - Trang "
        footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Đặt phong cách chung mặc định cho văn bản (Normal Style)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(6)

    # Tiêu đề Header văn bản hành chính Quốc hiệu
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_h = p_header.add_run("BỘ GIÁO DỤC VÀ ĐÀO TẠO - HỆ THỐNG TRỢ LÝ GIÁO ÁN SỐ\n")
    run_h.bold = True
    run_h.font.size = Pt(12)
    
    # Tên Giáo án chính
    info = lesson_json.get("thong_tin_chung", {})
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(f"KẾ HOẠCH BÀI DẠY PHÁT TRIỂN NĂNG LỰC SỐ\nTÊN BÀI HỌC: {info.get('ten_bai_hoc', '').upper()}\n")
    run_title.bold = True
    run_title.font.size = Pt(15)
    
    doc.add_paragraph(f"Môn học: {info.get('mon_hoc')}   |   Lớp: {info.get('lop')}   |   Thời lượng: {info.get('thoi_luong')}")
    
    # Mục I: Mục tiêu
    h1 = doc.add_paragraph()
    h1.add_run("I. MỤC TIÊU BÀI HỌC").bold = True
    
    doc.add_paragraph().add_run("1. Phẩm chất").bold = True
    for pc in lesson_json.get("muc_tieu", {}).get("pham_chat", []):
        doc.add_paragraph(f"- {pc}", style='List Bullet')
        
    doc.add_paragraph().add_run("2. Năng lực chung").bold = True
    for nlc in lesson_json.get("muc_tieu", {}).get("nang_luc_chung", []):
        doc.add_paragraph(f"- {nlc}", style='List Bullet')
        
    doc.add_paragraph().add_run("3. Năng lực đặc thù").bold = True
    for nldt in lesson_json.get("muc_tieu", {}).get("nang_luc_dac_thu", []):
        doc.add_paragraph(f"- {nldt}", style='List Bullet')

    # Mục II: Thiết bị dạy học
    h2 = doc.add_paragraph()
    h2.add_run("II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU SỐ").bold = True
    for tb in lesson_json.get("thiet_bi_day_hoc", []):
        doc.add_paragraph(f"- {tb}", style='List Bullet')

    # Mục III: Tiến trình dạy học Công văn 5512
    h3 = doc.add_paragraph()
    h3.add_run("III. TIẾN TRÌNH DẠY HỌC (CÔNG VĂN 5512)").bold = True
    
    activities = lesson_json.get("tien_trinh_day_hoc", [])

    if "Loại thường" in lesson_format:
        # --- LOẠI THƯỜNG (KHÔNG CHIA CỘT) ---
        for hd in activities:
            p_hd = doc.add_paragraph()
            p_hd.add_run(hd.get("ten_hoat_dong", "Hoạt động")).bold = True
            p_hd.paragraph_format.space_before = Pt(12)
            
            doc.add_paragraph().add_run("a) Mục tiêu: ").bold = True
            doc.add_paragraph(hd.get("muc_tieu", ""))
            
            doc.add_paragraph().add_run("b) Nội dung: ").bold = True
            doc.add_paragraph(hd.get("noi_dung", ""))
            
            doc.add_paragraph().add_run("c) Sản phẩm: ").bold = True
            doc.add_paragraph(hd.get("san_pham", ""))
            
            doc.add_paragraph().add_run("d) Tổ chức thực hiện: ").bold = True
            doc.add_paragraph(hd.get("to_chuc_thuc_hien", ""))

    elif "2 cột" in lesson_format:
        # --- LOẠI 2 CỘT ---
        for hd in activities:
            p_hd = doc.add_paragraph()
            p_hd.add_run(hd.get("ten_hoat_dong", "Hoạt động")).bold = True
            p_hd.paragraph_format.space_before = Pt(12)
            
            doc.add_paragraph().add_run("a) Mục tiêu: ").bold = True
            doc.add_paragraph(hd.get("muc_tieu", ""))
            
            doc.add_paragraph().add_run("b) Sản phẩm: ").bold = True
            doc.add_paragraph(hd.get("san_pham", ""))

            doc.add_paragraph().add_run("c) Nội dung & Tổ chức thực hiện: ").bold = True
            
            # Tạo bảng 2 cột: [Hoạt động của GV & HS] | [Nội dung cần đạt]
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Tiêu đề bảng
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Hoạt động của Giáo viên & Học sinh"
            hdr_cells[1].text = "Nội dung kiến thức cần đạt"
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_margins(cell)
            
            # Thêm dòng nội dung hoạt động
            row_cells = table.add_row().cells
            row_cells[0].text = hd.get("to_chuc_thuc_hien", "")
            row_cells[1].text = hd.get("noi_dung", "")
            for cell in row_cells:
                set_cell_margins(cell)
            doc.add_paragraph() # Cách dòng sau mỗi bảng

    elif "3 cột" in lesson_format:
        # --- LOẠI 3 CỘT ---
        for hd in activities:
            p_hd = doc.add_paragraph()
            p_hd.add_run(hd.get("ten_hoat_dong", "Hoạt động")).bold = True
            p_hd.paragraph_format.space_before = Pt(12)
            
            doc.add_paragraph().add_run("a) Mục tiêu: ").bold = True
            doc.add_paragraph(hd.get("muc_tieu", ""))
            
            doc.add_paragraph().add_run("b) Sản phẩm: ").bold = True
            doc.add_paragraph(hd.get("san_pham", ""))

            doc.add_paragraph().add_run("c) Tiến trình chi tiết: ").bold = True
            
            # Tạo bảng 3 cột: [Tiến trình] | [Hoạt động của GV & HS] | [Nội dung cần đạt]
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Tiến trình / Bước thực hiện"
            hdr_cells[1].text = "Hoạt động của Giáo viên & Học sinh"
            hdr_cells[2].text = "Nội dung cần đạt"
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_margins(cell)
            
            # Thêm dòng nội dung hoạt động
            row_cells = table.add_row().cells
            row_cells[0].text = "Chuyển giao nhiệm vụ -> Thực hiện -> Báo cáo thảo luận -> Đánh giá nhận xét."
            row_cells[1].text = hd.get("to_chuc_thuc_hien", "")
            row_cells[2].text = hd.get("noi_dung", "")
            for cell in row_cells:
                set_cell_margins(cell)
            doc.add_paragraph() # Cách dòng sau mỗi bảng

    # Mục IV: Xuất bảng biểu ma trận mục tiêu bài dạy
    doc.add_page_break()
    h4 = doc.add_paragraph()
    h4.add_run("IV. MA TRẬN MỤC TIÊU BÀI DẠY 4.0").bold = True
    
    t_matrix = doc.add_table(rows=1, cols=len(matrix_df.columns))
    t_matrix.style = 'Table Grid'
    
    hdr_cells = t_matrix.rows[0].cells
    for i, col_name in enumerate(matrix_df.columns):
        hdr_cells[i].text = col_name
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        set_cell_margins(hdr_cells[i])
        
    for index, row in matrix_df.iterrows():
        row_cells = t_matrix.add_row().cells
        for i in range(len(matrix_df.columns)):
            row_cells[i].text = str(row.iloc[i])
            set_cell_margins(row_cells[i])

    # Mục V: Phiếu học tập đi kèm bài dạy
    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    h5 = doc.add_paragraph()
    h5.add_run("V. PHỤ LỤC: PHIẾU HỌC TẬP TƯƠNG TÁC SỐ").bold = True
    doc.add_paragraph(worksheet_data.get("title", "")).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph().add_run("[Phiếu cá nhân số]").bold = True
    for item in worksheet_data.get("phieu_ca_nhan", []):
        doc.add_paragraph(item)
        
    doc.add_paragraph().add_run("[Phiếu thảo luận nhóm số]").bold = True
    for item in worksheet_data.get("phieu_nhom", []):
        doc.add_paragraph(item)

    # Ghi file vào bộ nhớ đệm
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
