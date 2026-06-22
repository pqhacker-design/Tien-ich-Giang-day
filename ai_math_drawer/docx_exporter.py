import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocxExporter:
    def __init__(self):
        self.doc = Document()
        # Thiết lập font chữ mặc định cho tài liệu
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

    def add_math_problem(self, index: int, title: str, img_stream: io.BytesIO):
        """Thêm một đề bài và hình vẽ tương ứng vào file Word"""
        # Thêm tiêu đề đề bài
        p_title = self.doc.add_paragraph()
        run_title = p_title.add_run(f"Bài toán {index}: {title}")
        run_title.bold = True
        
        # Thêm hình vẽ (Căn giữa)
        p_img = self.doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(img_stream, width=Inches(5.0))
        
        # Thêm chú thích dưới hình
        p_caption = self.doc.add_paragraph()
        p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = p_caption.add_run(f"Hình {index}. Hình vẽ minh họa cho Bài toán {index}")
        run_cap.italic = True
        run_cap.font.size = Pt(10)
        
        # Thêm khoảng cách dòng
        self.doc.add_paragraph()

    def save(self) -> io.BytesIO:
        """Lưu file Word vào bộ nhớ tạm để Streamlit download"""
        target_stream = io.BytesIO()
        self.doc.save(target_stream)
        target_stream.seek(0)
        return target_stream
