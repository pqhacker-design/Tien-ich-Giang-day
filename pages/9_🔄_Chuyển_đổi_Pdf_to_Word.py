import streamlit as st
import io
import re
from google import genai
from google.genai import types
import pdfplumber
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CẤU HÌNH GIAO DIỆN ---
st.set_page_config(page_title="Chuyển đổi Tài liệu sang Word", page_icon="🔄", layout="centered")

st.title("🔄 Chuyển đổi Tài liệu sang Word (Giữ nguyên Định dạng gốc)")
st.caption("Phiên bản Cao cấp: Giữ chính xác in đậm, in nghiêng, thụt lề, bảng biểu và đưa công thức về Unicode thuần.")

# --- LẤY API KEY TẬP TRUNG TỪ TRANG CHỦ ---
if "gemini_api_key" in st.session_state and st.session_state["gemini_api_key"].strip() != "":
    api_key_input = st.session_state["gemini_api_key"]
else:
    st.warning("⚠️ Vui lòng quay lại **Trang chủ** để nhập Google Gemini API Key trước khi sử dụng tính năng này.")
    st.stop() 

if "gemini_client" not in st.session_state:
    try:
        st.session_state.gemini_client = genai.Client(api_key=api_key_input)
    except Exception as e:
        st.error(f"Khởi tạo Gemini Client thất bại: {e}")
        st.stop()

client = st.session_state.gemini_client

# ==================== BỘ PHÂN TÍCH CHỮ IN ĐẬM / IN NGHIÊNG TỪ AI ====================
# ==================== BỘ PHÂN TÍCH CHỮ IN ĐẬM / IN NGHIÊNG TỪ AI ====================
def add_formatted_text(paragraph, text):
    """
    Hàm tự động quét qua chuỗi văn bản chứa ký tự markdown (* hoặc **)
    để phân tách, áp dụng thuộc tính Bold, Italic và XÓA BỎ dấu $ của công thức.
    """
    # Tách chuỗi dựa trên các token định dạng ẩn của markdown
    tokens = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    
    for token in tokens:
        if token.startswith('***') and token.endswith('***'):
            # Vừa đậm vừa nghiêng - Làm sạch dấu $
            clean_token = token[3:-3].replace('$', '')
            run = paragraph.add_run(clean_token)
            run.bold = True
            run.italic = True
        elif token.startswith('**') and token.endswith('**'):
            # Chữ in đậm - Làm sạch dấu $
            clean_token = token[2:-2].replace('$', '')
            run = paragraph.add_run(clean_token)
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            # Chữ in nghiêng - Làm sạch dấu $
            clean_token = token[1:-1].replace('$', '')
            run = paragraph.add_run(clean_token)
            run.italic = True
        # (Tìm đến đoạn chữ thường trong hàm add_formatted_text)
        else:
            # Chữ thường - Làm sạch dấu $ và các ký tự LaTeX rác nếu AI lọt lưới
            clean_token = token.replace('$', '')
            clean_token = clean_token.replace('\\cdot', '·')
            clean_token = clean_token.replace('\\neq', '≠')
            clean_token = clean_token.replace('\\leq', '≤')
            clean_token = clean_token.replace('\\geq', '≥')
            paragraph.add_run(clean_token)

# ==================== HÀM TẠO FILE WORD GIỮ ĐỊNH DẠNG TỐI ĐA ====================
def create_word_document(text_content):
    doc = Document()
    
    # Định dạng căn lề trang chuẩn văn bản hành chính Việt Nam (Lề 2cm mỗi bên)
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.79)

    # Đặt font chữ mặc định cho toàn văn bản là Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    lines = text_content.split('\n')
    in_table = False
    table_data = []

    for line in lines:
        raw_line = line
        clean_line = line.strip()
        if not clean_line:
            continue

        # 1. DỰNG LẠI BẢNG BIỂU CHUẨN Ô LƯỚI
        if clean_line.startswith('|'):
            if '---|' in clean_line or '===' in clean_line:
                continue
            in_table = True
            columns = [col.strip() for col in clean_line.split('|')[1:-1]]
            if columns:
                table_data.append(columns)
            continue
        else:
            if in_table and table_data:
                num_rows = len(table_data)
                num_cols = max(len(row) for row in table_data)
                word_table = doc.add_table(rows=num_rows, cols=num_cols)
                word_table.style = 'Table Grid'
                
                for r_idx, row in enumerate(table_data):
                    for c_idx, val in enumerate(row):
                        if c_idx < len(word_table.rows[r_idx].cells):
                            p_cell = word_table.rows[r_idx].cells[c_idx].paragraphs[0]
                            add_formatted_text(p_cell, val)
                
                doc.add_paragraph()
                in_table = False
                table_data = []

        # 2. XỬ LÝ ĐÁP ÁN TRẮC NGHIỆM CĂN NGANG HÀNG
        if re.search(r'A\..*B\..*C\..*D\.', clean_line) or re.search(r'A\s*[\.\)].*B\s*[\.\)].*C\s*[\.\)].*D\s*[\.\)]', clean_line):
            p = doc.add_paragraph()
            # Bảo toàn thụt lề nếu đầu dòng có khoảng trống
            if raw_line.startswith('    ') or raw_line.startswith('\t'):
                p.paragraph_format.left_indent = Inches(0.4)
                
            parts = re.split(r'(A\.|B\.|C\.|D\.)', clean_line)
            for part in parts:
                if part in ['A.', 'B.', 'C.', 'D.']:
                    run = p.add_run("    " + part + " ")
                    run.bold = True
                else:
                    add_formatted_text(p, part)
            continue

        # 3. XỬ LÝ TIÊU ĐỀ HOẶC CÁC CÂU HỎI VÀ VĂN BẢN THƯỜNG
        if clean_line.startswith('# '):
            h = doc.add_heading('', level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_text(h, clean_line.replace('# ', ''))
        elif clean_line.startswith('## '):
            h = doc.add_heading('', level=2)
            add_formatted_text(h, clean_line.replace('## ', ''))
        else:
            p = doc.add_paragraph()
            
            # --- XỬ LÝ THỤT LỀ ĐẦU DÒNG (INDENTATION) ---
            if raw_line.startswith('    ') or raw_line.startswith('\t'):
                p.paragraph_format.left_indent = Inches(0.4)

            if clean_line.startswith("Câu "):
                match = re.match(r'(Câu \d+\.)(.*)', clean_line)
                if match:
                    bold_part, regular_part = match.groups()
                    p.add_run(bold_part).bold = True
                    add_formatted_text(p, regular_part)
                else:
                    add_formatted_text(p, clean_line)
            else:
                add_formatted_text(p, clean_line)

    # Đề phòng bảng biểu nằm ở dòng cuối cùng của file
    if in_table and table_data:
        word_table = doc.add_table(rows=len(table_data), cols=max(len(row) for row in table_data))
        word_table.style = 'Table Grid'
        for r_idx, row in enumerate(table_data):
            for c_idx, val in enumerate(row):
                if c_idx < len(word_table.rows[r_idx].cells):
                    p_cell = word_table.rows[r_idx].cells[c_idx].paragraphs[0]
                    add_formatted_text(p_cell, val)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- GIAO DIỆN STREAMLIT ---
uploaded_file = st.file_uploader(
    "Tải lên Tài liệu dạng Ảnh hoặc PDF cần giữ nguyên định dạng:",
    type=["png", "jpg", "jpeg", "pdf"]
)

if uploaded_file is not None:
    file_type = uploaded_file.type
    extracted_text = ""
    
    st.info(f"📁 Đã nhận file: **{uploaded_file.name}**")
    
    if st.button("🚀 Bắt đầu chuyển đổi bảo toàn định dạng"):
        with st.spinner("Hệ thống AI đang phân tích cấu trúc chữ đậm, chữ nghiêng và thụt lề gốc..."):
            
            if file_type == "application/pdf":
                try:
                    with pdfplumber.open(uploaded_file) as pdf:
                        text_pages = []
                        for page in pdf.pages:
                            page_text = page.extract_text(layout=True)
                            if page_text:
                                text_pages.append(page_text)
                        extracted_text = "\n".join(text_pages)
                except Exception:
                    pass
            
            if not extracted_text.strip():
                uploaded_file.seek(0)
                file_bytes = uploaded_file.read()
                file_part = types.Part.from_bytes(data=file_bytes, mime_type=file_type)
                
                # Prompt ÉP AI phân biệt chữ Đậm, chữ Nghiêng bằng cú pháp Markdown cụ thể
                # Bộ Prompt nâng cấp chuyên sâu ép AI chuyển đổi Toán học sang Unicode phẳng trực quan
                prompt = (
                    "Bạn là một chuyên gia số hóa tài liệu giáo dục và đề thi toán học cấp cao.\n"
                    "Hãy trích xuất chính xác 100% văn bản từ hình ảnh này sang định dạng Unicode phẳng.\n\n"
                    
                    "🔴 QUY TẮC CHUYỂN ĐỔI TOÁN HỌC SANG UNICODE TRỰC QUAN (CẤM DÙNG LATEX HOẶC DẤU $):\n"
                    "1. Phân số phức tạp: Hãy viết rõ ràng theo hàng ngang trực quan bằng ký tự Unicode, hoặc viết dạng tử/mẫu cách khoảng rõ ràng, ví dụ: 'f_1(x) = (5x⁴)/4' hoặc '(x + 1)/(x - 2)'. Tránh viết dính liền gây hiểu lầm.\n"
                    "2. Hệ phương trình (cases): Chuyển thành dạng ký tự phẳng sử dụng dấu ngoặc nhọn lớn '{' và xuống dòng rõ ràng cho từng phương trình, ví dụ:\n"
                    "   Hệ phương trình: {\n"
                    "   x + y - 2 < 0\n"
                    "   x - y + 2 > 0\n"
                    "3. Số mũ và Chỉ số: BẮT BUỘC dùng ký tự Unicode nhỏ trên và dưới (Ví dụ: x², y³, u_n, u₁, u₂, log₃(3x), f'(x)). Không được để dạng dấu mũ ^ hoặc gạch dưới _ thô.\n"
                    "4. Ký hiệu hình học & Giải tích: Chuyển toàn bộ về ký tự Unicode tương ứng:\n"
                    "   - Vectơ: Dùng mũi tên đứng trước (Ví dụ: →AD, →AB).\n"
                    "   - Tích phân / Nguyên hàm: Dùng ký tự ∫ (Ví dụ: ∫f(x)dx).\n"
                    "   - Tập hợp & Logic: Dùng chuẩn ký tự ℝ, ∈, ∉, ⊂, ∩, ∪, ≠, ≥, ≤, · (dấu nhân).\n\n"
                    
                    "🔴 QUY TẮC ĐỊNH DẠNG VĂN BẢN ĐỀ THI:\n"
                    "1. BẢO TOÀN CHỮ IN ĐẬM/IN NGHIÊNG: Đoạn văn nào in đậm bắt buộc bọc trong `**` (Ví dụ: **BỘ GIÁO DỤC VÀ ĐÀO TẠO**), đoạn nào in nghiêng bọc trong `*` (Ví dụ: *xem hình dưới*).\n"
                    "2. CĂN NGANG ĐÁP ÁN: 4 đáp án trắc nghiệm A, B, C, D phải nằm ngang trên cùng một dòng văn bản giống hệt đề gốc, cách nhau bằng khoảng trắng lớn.\n"
                    "3. THỤT LỀ ĐẦU DÒNG: Thêm 4 dấu cách '    ' vào trước mỗi dòng câu hỏi (Câu 1, Câu 2...) hoặc đoạn văn có thụt lề.\n"
                    "4. Chỉ trả về nội dung đề thi đã số hóa sạch sẽ, không thêm bất kỳ lời thoại hay giải thích nào của AI."
                )
                
                try:
                    response = client.models.generate_content(
                        model="gemini-3.5-flash",
                        contents=[file_part, prompt]
                    )
                    extracted_text = response.text
                except Exception as e:
                    st.error(f"Lỗi khi kết nối API: {e}")

        if extracted_text.strip():
            st.success("✅ Đã chuẩn hóa định dạng thành công!")
            
            with st.expander("👀 Xem trước văn bản trích xuất (Đã gán tag định dạng)"):
                st.text_area("Mã nguồn văn bản:", extracted_text, height=300)
            
            word_file = create_word_document(extracted_text)
            
            st.download_button(
                label="📥 Tải file WORD (.docx) Full Định Dạng về máy",
                data=word_file,
                file_name=uploaded_file.name.rsplit('.', 1)[0] + "_full_dinh_dang.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning("Không thể xử lý cấu trúc từ file.")

# --- FOOTER CỐ ĐỊNH ---
st.divider()
st.markdown(
    """
    <div style="text-align: center; font-size: 0.8em; color: grey;">
        Ứng dụng được phát triển bởi Ngô Thanh Hùng
    </div>
    """,
    unsafe_allow_html=True
)
