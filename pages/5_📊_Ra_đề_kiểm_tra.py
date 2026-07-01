import streamlit as st
import json
import os
import re
import random
import copy
import zipfile
from io import BytesIO
import pandas as pd

# Thư viện xử lý tài liệu Word chuyên sâu
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

import google.generativeai as genai

# ==========================================
# CẤU HÌNH TRANG & GIAO DIỆN
# ==========================================
st.set_page_config(
    page_title="SmartTest MultiSport - Hệ thống Tạo Đề Đa Môn Toàn Diện",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
/* Nhắm mục tiêu vào văn bản bên trong các nút Tab của Streamlit */
        button[data-baseweb="tab"] div p {
            font-weight: bold !important;
            font-size: 1.05em !important; /* Có thể phóng to chữ lên một chút nếu muốn */
        }
    .main-title { font-size: 2.3rem; font-weight: 700; color: #1E3A8A; text-align: center; margin-bottom: 1.5rem; }
    .section-header { font-size: 1.3rem; font-weight: 600; color: #0F766E; margin-top: 1.5rem; margin-bottom: 1rem; border-left: 5px solid #0F766E; padding-left: 10px; }
    .stButton>button { width: 100%; background-color: #0F766E; color: white; border-radius: 8px; font-weight: 600; padding: 10px; }
    .stButton>button:hover { background-color: #0D9488; color: white; }
    .score-box { background-color: #F0FDF4; padding: 15px; border-radius: 8px; border: 1px solid #BBF7D0; margin-bottom: 10px; }
</style>
""", unsafe_allow_html=True)

# ==========================================
# DANH SÁCH TẤT CẢ CÁC MÔN HỌC & ĐẶC THÙ KÝ HIỆU
# ==========================================
SUBJECTS_CONFIG = {
    "Toán học": "Sử dụng ký tự Unicode toán học trực tiếp trong văn bản trơn. Số mũ dùng kí tự mũ trực tiếp như ², ³, ⁴ (x², cm²). Phân số viết dạng chia ngang (x + 1)/(x - 3). Hình học dùng ΔABC, ∽, ⊥, //.",
    "Vật lý": "Đơn vị đo lường phải chuẩn hóa (Ω, W, J, m/s²). Kí hiệu mũ trực tiếp như s², m³.",
    "Hóa học": "BẮT BUỘC dùng chỉ số dưới Unicode thực tế cho công thức phân tử: H₂SO₄, Ca(OH)₂, Al₂(SO₄)₃. Mũi tên phản ứng dạng → hoặc ⇌.",
    "Tin học": "Các đoạn mã giả, code Python, C++, HTML hoặc ký hiệu logic (AND, OR, NOT, ⊕) phải đặt trong dấu `...` hoặc bọc khối rõ ràng.",
    "Công nghệ": "Ký hiệu mạch điện, thông số kỹ thuật điện trở, bản vẽ kỹ thuật hoặc sơ đồ khối phải mô tả tường minh bằng ký hiệu hoa văn chuẩn.",
    "Ngữ văn": "Các ngữ liệu văn học, đoạn thơ, đoạn văn trích dẫn phải được bọc trong dấu ngoặc kép hoặc ghi rõ nguồn. Câu hỏi trắc nghiệm/tự luận tập trung vào đọc hiểu và làm văn.",
    "Tiếng Anh / Ngoại ngữ": "Toàn bộ câu hỏi, phương án lựa chọn và văn bản đọc hiểu phải viết bằng ngôn ngữ đích chuẩn bản xứ, không sai chính tả, có phần trọng âm, ngữ âm rõ ràng.",
    "Lịch sử & Địa lý": "Mốc thời gian, số liệu thống kê tọa độ, ký hiệu bản đồ, số liệu kinh tế - xã hội phải chính xác tuyệt đối theo dòng sự kiện và Atlat.",
    "Giáo dục Kinh tế và Pháp luật": "Các thuật ngữ pháp lý, điều luật, tình huống thực tế phải bọc trong quy chuẩn trích dẫn lập pháp.",
    "Sinh học / KHTN": "Ký hiệu sơ đồ phép lai (P, F₁, F₂), alen (A₁, a), nhiễm sắc thể (2n, n) hoặc công thức phân tử sinh học phải viết chuẩn xác."
}

# ==========================================
# TRÍCH XUẤT VÀ KẾT NỐI GEMINI AI
# ==========================================
def extract_text_from_docx(file_bytes):
    try:
        doc = Document(BytesIO(file_bytes))
        fullText = []
        for para in doc.paragraphs:
            if para.text.strip():
                fullText.append(para.text)
        return "\n".join(fullText)
    except Exception as e:
        return f"Lỗi đọc file Word: {str(e)}"

def add_math_run_to_paragraph(paragraph, text):
    if not text: return
    text = str(text).replace('\\n', '\n').replace('\\', '').replace('$', '').strip()
    lines = text.split('\n')
    for index, line in enumerate(lines):
        if index > 0:
            new_run = paragraph.add_run()
            new_run.add_break()
        if line.strip():
            paragraph.add_run(line)

# ==========================================
# KHỞI TẠO SESSION STATE
# ==========================================
if 'generated_data' not in st.session_state: st.session_state.generated_data = None
if 'step1_data' not in st.session_state: st.session_state.step1_data = None
if 'multi_codes_data' not in st.session_state: st.session_state.multi_codes_data = {}
if 'alignment_table' not in st.session_state: st.session_state.alignment_table = None
if 'user_custom_req' not in st.session_state: st.session_state.user_custom_req = ""

# ==========================================
# THUẬT TOÁN ĐẢO ĐỀ MULTI-CODE
# ==========================================
def generate_shuffled_bundle(original_data, start_code, num_codes):
    bundle = {}
    alignment_records = []
    de_goc = original_data.get('de_kiem_tra', {})
    tn_4_lua_chon_goc = de_goc.get('trac_nghiem_4_lua_chon', [])
    
    if not tn_4_lua_chon_goc and not de_goc.get('trac_nghiem_dung_sai') and not de_goc.get('trac_nghiem_tra_loi_ngan'):
        return bundle, None

    for i in range(num_codes):
        current_code = str(start_code + i)
        shuffled_data = copy.deepcopy(original_data)
        
        if tn_4_lua_chon_goc:
            tn_current = shuffled_data['de_kiem_tra']['trac_nghiem_4_lua_chon']
            random.seed(int(current_code) + 200)
            indexed_tn = list(enumerate(tn_current))
            random.shuffle(indexed_tn)
            
            new_tn_list = []
            new_dap_an_tn = {}
            
            for new_idx, (old_idx, q) in enumerate(indexed_tn):
                new_id = f"{new_idx + 1}"
                old_id = f"Câu {old_idx + 1}"
                
                q['id'] = int(new_id)
                # Thuộc tính q['muc_do'] được giữ nguyên nhờ deepcopy
                
                opts = { "A": q.get("A",""), "B": q.get("B",""), "C": q.get("C",""), "D": q.get("D","") }
                old_correct_key = shuffled_data['dap_an_chi_tiet']['trac_nghiem_4_lua_chon'].get(str(old_idx + 1))
                old_correct_value = opts.get(old_correct_key)
                
                opt_values = list(opts.values())
                random.shuffle(opt_values)
                
                for o_idx, char in enumerate(['A', 'B', 'C', 'D']):
                    q[char] = opt_values[o_idx]
                    if opt_values[o_idx] == old_correct_value:
                        new_correct_key = char
                
                new_dap_an_tn[str(new_id)] = new_correct_key
                new_tn_list.append(q)
                
                alignment_records.append({
                    "Mã đề": current_code, "Câu hỏi gốc": old_id, "Vị trí mới": f"Câu {new_id}", "Đáp án mới": new_correct_key
                })
                
            shuffled_data['de_kiem_tra']['trac_nghiem_4_lua_chon'] = new_tn_list
            shuffled_data['dap_an_chi_tiet']['trac_nghiem_4_lua_chon'] = new_dap_an_tn
            
        bundle[current_code] = shuffled_data

    if alignment_records:
        df_log = pd.DataFrame(alignment_records)
        pivot_df = df_log.pivot(index='Câu hỏi gốc', columns='Mã đề', values='Vị trí mới')
        pivot_df = pivot_df.reindex(index=sorted(pivot_df.index, key=lambda x: int(re.search(r'\d+', x).group()))).reset_index()
        return bundle, pivot_df
    return bundle, None

def init_gemini_client(api_key):
    try:
        genai.configure(api_key=api_key)
        return genai.GenerativeModel('gemini-2.5-flash')
    except Exception as e:
        st.error(f"Lỗi API Key hoặc cấu hình kết nối: {e}")
        return None

def generate_step1_matrix(model, config, raw_input_data):
    prompt_text = f"""
    Bạn là chuyên gia khảo thí cấp cao của Bộ Giáo dục. Hãy lập Ma trận và Bản đặc tả đề kiểm tra môn {config['subject']} - Khối Lớp {config['grade']}.
    Độ khó mục tiêu: {config['difficulty']}.
    
    ĐẶC BIỆT LƯU Ý VỀ THỜI GIAN LÀM BÀI VÀ ĐỐI TƯỢNG:
    - Thời lượng làm bài: {config['duration']} phút.
    - Độ khó mục tiêu: {config['difficulty']}.
    
    YÊU CẦU ĐIỀU CHỈNH ĐỘ KHÓ THEO THỜI LƯỢNG:
    - Nếu thời lượng NGẮN (≤ 15-20 phút): Các câu hỏi thuộc mức độ Thông hiểu và Vận dụng phải là các bài toán/tình huống xử lý NHANH (tối đa 1-2 bước tính), tránh các câu hỏi cần đọc hiểu ngữ liệu quá dài hoặc tính toán phức tạp để học sinh kịp làm bài.
    - Nếu thời lượng DÀI (≥ 45 phút): Bản đặc tả được phép mở rộng các tiêu chí đánh giá, yêu cầu các câu hỏi Vận dụng và Vận dụng cao có độ sâu tư duy, bài toán nhiều bước giải, hoặc đoạn văn/ngữ liệu phân tích dài để phân hóa tốt đối tượng học sinh tương ứng với mức độ ({config['difficulty']}).
    Số câu Trắc nghiệm 4 lựa chọn: {config['num_tn_4_lua_chon']} ({config['score_part1']} điểm), 
    Số câu Đúng/Sai: {config['num_tn_dung_sai']} ({config['score_part2']} điểm), 
    Số câu Trả lời ngắn: {config['num_tn_tra_loi_ngan']} ({config['score_part3']} điểm), 
    Số câu Tự luận: {config['num_tl']} ({config['score_part4']} điểm).
    Đặc biệt, lưu ý phân phối điểm số riêng cho câu hỏi Vận dụng cao (VDC) trong đề là: {config['score_vdc_custom']} điểm.
    Tổng điểm của toàn đề bắt buộc là 10 điểm.
    Tỷ lệ nhận thức: Nhận biết {config['nb_ratio']}% , Thông hiểu {config['th_ratio']}%, Vận dụng {config['vd_ratio']}%, Vận dụng cao {config['vdc_ratio']}%.
    Mẫu thiết kế ma trận yêu cầu: "{config['matrix_template']}".

    YÊU CẦU BẮT BUỘC:
    - Phân tích thật sâu tài liệu/hình ảnh/yêu cầu đính kèm đi cùng lệnh này để phân chia thành các "chu_de" và "noi_dung".
    - Phân chia chi tiết số câu hỏi tương ứng dựa trên tổng điểm và số câu của từng phần đã định cấu hình.

    BẮT BUỘC TRẢ VỀ JSON NGUYÊN BẢN CÓ CẤU TRÚC SAU:
    {{
      "ma_tran": [
        {{
          "tt": 1,
          "chu_de": "Tên chủ đề/Chương lớn",
          "noi_dung": "Nội dung kiến thức cụ thể",
          "nhieu_lua_chon": {{"nb": 2, "th": 1, "vd": 0}},
          "dung_sai": {{"nb": 0, "th": 1, "vd": 0}},
          "tra_loi_ngan": {{"nb": 1, "th": 0}},
          "tu_luan": {{"nb": 0, "th": 0, "vd": 0, "vdc": 0}},
          "tong_diem_phan_tram": 15
        }}
      ],
      "bang_dac_ta": [
        {{"id_dac_ta": "DT_01", "chu_de": "Tên chủ đề", "noi_dung": "Kiến thức cụ thể", "muc_do": "Nhận biết", "yeu_cau_can_dat": "Mô tả yêu cầu cần đạt...", "so_cau": "2 câu TN", "diem": 0.5}}
      ]
    }}
    """
    if st.session_state.get("user_custom_req"):
        prompt_text += f"\n[Yêu cầu tùy biến bổ sung từ giáo viên: {st.session_state.user_custom_req}]\n"

    contents_to_send = []
    if isinstance(raw_input_data, dict) and "mime_type" in raw_input_data:
        contents_to_send.append(raw_input_data)
        contents_to_send.append(prompt_text)
    else:
        prompt_text += f"\nNGUỒN DỮ LIỆU KIẾN THỨC BẮT BUỘC:\n\"\"\"\n{raw_input_data}\n\"\"\""
        contents_to_send.append(prompt_text)

    response = model.generate_content(
        contents_to_send,
        generation_config={"response_mime_type": "application/json"}
    )
    raw = response.text.strip()
    return json.loads(raw, strict=False)

def generate_step2_questions(model, config, matrix_data, subject_rule, raw_input_data):
    dac_ta_str = json.dumps(matrix_data.get('bang_dac_ta', []), ensure_ascii=False)
    prompt_text = f"""
    Bạn là chuyên gia soạn đề thi bám sát cấu trúc đề minh họa mới của Bộ Giáo dục. Dựa trên Bản đặc tả sau: {dac_ta_str}
    Hãy thiết lập chi tiết nội dung đề kiểm tra môn {config['subject']}, Khối {config['grade']}.
    Đặc thù môn học: {subject_rule}

    THÔNG TIN THỜI GIAN LÀM BÀI BẮT BUỘC TUÂN THỦ:
    - Tổng thời gian làm bài: {config['duration']} phút cho tổng số {config['num_tn_4_lua_chon'] + config['num_tn_dung_sai'] + config['num_tn_tra_loi_ngan'] + config['num_tl']} câu hỏi.
    - Do đó, bạn phải tự cân đối và khống chế "độ dài" và "độ phức tạp" của từng câu hỏi sao cho một học sinh thuộc nhóm đối tượng ({config['difficulty']}) có thể hoàn thành bài thi một cách hợp lý trong đúng {config['duration']} phút.
    - VÍ DỤ: Đề {config['duration']} phút mà ngắn (ví dụ 15 phút) thì câu hỏi trắc nghiệm phải ngắn gọn, hỏi thẳng vào bản chất, công thức tính nhanh, phần tự luận (nếu có) phải là bài toán cơ bản. Không được ra câu hỏi tốn quá nhiều thời gian đọc hay tính toán.
    Yêu cầu số lượng câu hỏi và cơ cấu điểm số:
    - Phần I (Trắc nghiệm nhiều lựa chọn): {config['num_tn_4_lua_chon']} câu. Tổng điểm phần này: {config['score_part1']} điểm.
    - Phần II (Trắc nghiệm Đúng/Sai): {config['num_tn_dung_sai']} câu. Tổng điểm phần này: {config['score_part2']} điểm.
    - Phần III (Trắc nghiệm Trả lời ngắn): {config['num_tn_tra_loi_ngan']} câu. Tổng điểm phần này: {config['score_part3']} điểm.
    - Phần IV (Tự luận): {config['num_tl']} câu. Tổng điểm phần này: {config['score_part4']} điểm.
    
    * LƯU Ý PHÂN HÓA HỌC SINH: Câu hỏi thuộc mức độ Vận dụng cao (VDC) trong đề phải chiếm đúng trọng số tương đương {config['score_vdc_custom']} điểm trên tổng số điểm của phần chứa nó.

    QUY ĐỊNH ĐỊNH DẠNG TOÁN HỌC & KHTN (BẮT BUỘC KHÔNG DÙNG LATEX):
    - TUYỆT ĐỐI KHÔNG sử dụng ký tự $, \, frac, neq, sim, triangle, hoặc ^.
    - Sử dụng ký tự Unicode toán học trực tiếp trong văn bản trơn (Ví dụ: x², cm², ΔABC, ∽, ⊥, //, →, ⇌, ·, ≠, °).

    YÊU CẦU QUAN TRỌNG: Từng câu hỏi trong mỗi danh sách PHẢI được gán nhãn thuộc tính mức độ nhận thức "muc_do" (nhận giá trị là "NB", "TH", "VD", hoặc "VDC") dựa trên ma trận bản đặc tả.

    BẮT BUỘC TRẢ VỀ JSON NGUYÊN BẢN CÓ CẤU TRÚC SAU:
    {{
      "de_kiem_tra": {{
        "trac_nghiem_4_lua_chon": [
          {{"id": 1, "muc_do": "NB", "cau_hoi": "Nội dung câu hỏi...", "A": "Đáp án A", "B": "Đáp án B", "C": "Đáp án C", "D": "Đáp án D"}}
        ],
        "trac_nghiem_dung_sai": [
          {{"id": 1, "muc_do": "TH", "cau_hoi": "Nội dung câu hỏi...", "cac_y": {{"a": "Ý a", "b": "Ý b", "c": "Ý c", "d": "Ý d"}}}}
        ],
        "trac_nghiem_tra_loi_ngan": [
          {{"id": 1, "muc_do": "VD", "cau_hoi": "Nội dung câu hỏi..."}}
        ],
        "tu_luan": [
          {{"id": 1, "muc_do": "VDC", "cau_hoi": "Nội dung câu hỏi..."}}
        ]
      }},
      "dap_an_chi_tiet": {{
        "trac_nghiem_4_lua_chon": {{"1": "A"}},
        "trac_nghiem_dung_sai": {{"1": {{"a": "Đúng", "b": "Sai", "c": "Đúng", "d": "Sai"}}}},
        "trac_nghiem_tra_loi_ngan": {{"1": "25"}},
        "tu_luan": {{"1": "Lời giải..."}}
      }}
    }}
    """
    contents_to_send = []
    if isinstance(raw_input_data, dict) and "mime_type" in raw_input_data:
        contents_to_send.append(raw_input_data)
        contents_to_send.append(prompt_text)
    else:
        prompt_text += f"\nNGUỒN KIẾN THỨC BẮT BUỘC ĐỂ SOẠN CÂU HỎI:\n\"\"\"\n{raw_input_data}\n\"\"\""
        contents_to_send.append(prompt_text)

    response = model.generate_content(contents_to_send)
    raw = response.text.strip()
    raw = re.sub(r'^```json\s*|```$', '', raw, flags=re.IGNORECASE).strip()
    return json.loads(raw, strict=False)

def build_single_docx(config, data, code_label, include_matrix=True):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.79)
        
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    p_top = doc.add_paragraph()
    p_top.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_top.add_run("TRƯỜNG THCS & THPT Artificial Intelligence (AI)\n").bold = True
    
    exam_title = f"ĐỀ KIỂM TRA {config.get('exam_type', '').upper()} - MÃ ĐỀ: {code_label}\n"
    run_title = p_top.add_run(exam_title)
    run_title.bold = True
    run_title.size = Pt(14)
    
    p_top.add_run(f"Môn: {config['subject']} | Khối: {config['grade']} | Thời gian: {config['duration']} phút\n")
    p_top.add_run("-------------------------------------\n")
    
    if include_matrix and 'ma_tran' in data:
        doc.add_heading("I. MA TRẬN PHÂN BỔ ĐỀ KIỂM TRA", level=2)
        is_cv7991 = "7991" in config.get("matrix_template", "")
        
        if is_cv7991:
            m_table = doc.add_table(rows=1, cols=8)
            m_table.style = 'Table Grid'
            m_hdrs = ['TT', 'Chủ đề Chương', 'Nội dung kiến thức', 'Nhiều lựa chọn (NB/TH/VD)', 'Đúng - Sai (NB/TH/VD)', 'Trả lời ngắn (NB/TH)', 'Tự luận (NB/TH/VD/VDC)', 'Tổng %']
            for idx, h in enumerate(m_hdrs): 
                m_table.rows[0].cells[idx].text = h
                m_table.rows[0].cells[idx].paragraphs[0].runs[0].font.bold = True
                
            for item in data['ma_tran']:
                row = m_table.add_row().cells
                row[0].text = str(item.get('tt', '1'))
                row[1].text = str(item.get('chu_de', ''))
                row[2].text = str(item.get('noi_dung', ''))
                
                nlc = item.get('nhieu_lua_chon', {})
                ds = item.get('dung_sai', {})
                tln = item.get('tra_loi_ngan', {})
                tl = item.get('tu_luan', {})
                
                row[3].text = f"{nlc.get('nb',0)}/{nlc.get('th',0)}/{nlc.get('vd',0)}"
                row[4].text = f"{ds.get('nb',0)}/{ds.get('th',0)}/{ds.get('vd',0)}"
                row[5].text = f"{tln.get('nb',0)}/{tln.get('th',0)}"
                row[6].text = f"{tl.get('nb',0)}/{tl.get('th',0)}/{tl.get('vd',0)}/{tl.get('vdc',0)}"
                row[7].text = f"{item.get('tong_diem_phan_tram', 10)}%"
        else:
            # Mẫu đơn giản: Tách rõ số câu TN và TL theo từng mức độ nhận thức
            m_table = doc.add_table(rows=1, cols=5)
            m_table.style = 'Table Grid'
            m_hdrs = ['Chủ đề / Nội dung kiến thức', 'Nhận biết', 'Thông hiểu', 'Vận dụng', 'Vận dụng cao']
            for idx, h in enumerate(m_hdrs): 
                m_table.rows[0].cells[idx].text = h
                m_table.rows[0].cells[idx].paragraphs[0].runs[0].font.bold = True
            
            # Khởi tạo các biến tích lũy tổng số câu
            total_nb_tn, total_nb_tl = 0, 0
            total_th_tn, total_th_tl = 0, 0
            total_vd_tn, total_vd_tl = 0, 0
            total_vdc_tn, total_vdc_tl = 0, 0
            
            for item in data['ma_tran']:
                row = m_table.add_row().cells
                row[0].text = f"{item.get('chu_de', '')} - {item.get('noi_dung', '')}"
                
                nlc = item.get('nhieu_lua_chon', {})
                ds = item.get('dung_sai', {})
                tln = item.get('tra_loi_ngan', {})
                tl = item.get('tu_luan', {})
                
                # Phân nhóm: TN (Trắc nghiệm nhiều lựa chọn + Đúng sai + Trả lời ngắn) và TL (Tự luận)
                nb_tn = nlc.get('nb', 0) + ds.get('nb', 0) + tln.get('nb', 0)
                nb_tl = tl.get('nb', 0)
                
                th_tn = nlc.get('th', 0) + ds.get('th', 0) + tln.get('th', 0)
                th_tl = tl.get('th', 0)
                
                vd_tn = nlc.get('vd', 0) + ds.get('vd', 0)
                vd_tl = tl.get('vd', 0)
                
                vdc_tn = 0 
                vdc_tl = tl.get('vdc', 0)
                
                # Điền dữ liệu vào bảng
                row[1].text = f"TN: {nb_tn} | TL: {nb_tl}"
                row[2].text = f"TN: {th_tn} | TL: {th_tl}"
                row[3].text = f"TN: {vd_tn} | TL: {vd_tl}"
                row[4].text = f"TN: {vdc_tn} | TL: {vdc_tl}"
                
                # Tích lũy vào tổng
                total_nb_tn += nb_tn; total_nb_tl += nb_tl
                total_th_tn += th_tn; total_th_tl += th_tl
                total_vd_tn += vd_tn; total_vd_tl += vd_tl
                total_vdc_tn += vdc_tn; total_vdc_tl += vdc_tl

            # --- HÀNG TỔNG SỐ CÂU ---
            row_total_q = m_table.add_row().cells
            row_total_q[0].text = "Tổng số câu"
            row_total_q[1].text = f"TN: {total_nb_tn} | TL: {total_nb_tl}"
            row_total_q[2].text = f"TN: {total_th_tn} | TL: {total_th_tl}"
            row_total_q[3].text = f"TN: {total_vd_tn} | TL: {total_vd_tl}"
            row_total_q[4].text = f"TN: {total_vdc_tn} | TL: {total_vdc_tl}"
            
            # --- HÀNG TỔNG SỐ ĐIỂM ---
            row_total_p = m_table.add_row().cells
            row_total_p[0].text = "Tổng số điểm"
            
            row_total_p[1].text = f"{config.get('nb_ratio', 40) / 10} điểm"
            row_total_p[2].text = f"{config.get('th_ratio', 30) / 10} điểm"
            row_total_p[3].text = f"{config.get('vd_ratio', 20) / 10} điểm"
            row_total_p[4].text = f"{config.get('vdc_ratio', 10) / 10} điểm"
            
            for i in range(5):
                row_total_q[i].paragraphs[0].runs[0].font.bold = True
                row_total_p[i].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()
    # --- BỔ SUNG: VẼ BẢNG BẢN ĐẶC TẢ ---
        if 'bang_dac_ta' in data and data['bang_dac_ta']:
            doc.add_heading("II. BẢN ĐẶC TẢ PHƯƠNG ÁN RA ĐỀ", level=2)
            dt_table = doc.add_table(rows=1, cols=7)
            dt_table.style = 'Table Grid'
            dt_hdrs = ['Mã ĐT', 'Chủ đề', 'Nội dung', 'Mức độ', 'Yêu cầu cần đạt', 'Số câu', 'Điểm']
            
            # Ghi tiêu đề bảng
            for idx, h in enumerate(dt_hdrs):
                dt_table.rows[0].cells[idx].text = h
                dt_table.rows[0].cells[idx].paragraphs[0].runs[0].font.bold = True
                
            # Đổ dữ liệu từ JSON vào bảng Word
            for item in data['bang_dac_ta']:
                row = dt_table.add_row().cells
                row[0].text = str(item.get('id_dac_ta', ''))
                row[1].text = str(item.get('chu_de', ''))
                row[2].text = str(item.get('noi_dung', ''))
                row[3].text = str(item.get('muc_do', ''))
                row[4].text = str(item.get('yeu_cau_can_dat', ''))
                row[5].text = str(item.get('so_cau', ''))
                row[6].text = str(item.get('diem', ''))
            
            doc.add_paragraph() # Dòng trống giãn cách
            
    doc.add_heading("III. NỘI DUNG CÂU HỎI", level=2)
    de = data.get('de_kiem_tra', {})
    
    list_4lc = de.get("trac_nghiem_4_lua_chon", [])
    if list_4lc:
        p_hdr1 = doc.add_paragraph()
        p_hdr1.add_run(f"PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn ({config.get('score_part1', 3.0)} điểm).").bold = True
        for q in list_4lc:
            p_q = doc.add_paragraph()
            muc_do = q.get('muc_do', 'NB')
            p_q.add_run(f"Câu {q.get('id', '')} ({muc_do}_TN): ").bold = True
            add_math_run_to_paragraph(p_q, q.get('cau_hoi', ''))
            p_opts = doc.add_paragraph()
            p_opts.paragraph_format.left_indent = Inches(0.3)
            p_opts.add_run("A. ").bold = True
            add_math_run_to_paragraph(p_opts, q.get('A', ''))
            p_opts.add_run("   B. ").bold = True
            add_math_run_to_paragraph(p_opts, q.get('B', ''))
            p_opts.add_run("   C. ").bold = True
            add_math_run_to_paragraph(p_opts, q.get('C', ''))
            p_opts.add_run("   D. ").bold = True
            add_math_run_to_paragraph(p_opts, q.get('D', ''))

    list_ds = de.get("trac_nghiem_dung_sai", [])
    if list_ds:
        p_hdr2 = doc.add_paragraph()
        p_hdr2.add_run(f"\nPHẦN II. Câu trắc nghiệm đúng sai ({config.get('score_part2', 2.0)} điểm).").bold = True
        for q in list_ds:
            p_q = doc.add_paragraph()
            muc_do = q.get('muc_do', 'TH')
            p_q.add_run(f"Câu {q.get('id', '')} ({muc_do}_TN): ").bold = True
            add_math_run_to_paragraph(p_q, q.get('cau_hoi', ''))
            cac_y = q.get("cac_y", {})
            for key in ['a', 'b', 'c', 'd']:
                if key in cac_y:
                    p_y = doc.add_paragraph()
                    p_y.paragraph_format.left_indent = Inches(0.4)
                    p_y.add_run(f"{key}) ").bold = True
                    add_math_run_to_paragraph(p_y, cac_y.get(key, ''))

    list_tln = de.get("trac_nghiem_tra_loi_ngan", [])
    if list_tln:
        p_hdr3 = doc.add_paragraph()
        p_hdr3.add_run(f"\nPHẦN III. Câu trắc nghiệm trả lời ngắn ({config.get('score_part3', 2.0)} điểm).").bold = True
        for q in list_tln:
            p_q = doc.add_paragraph()
            muc_do = q.get('muc_do', 'VD')
            p_q.add_run(f"Câu {q.get('id', '')} ({muc_do}_TN): ").bold = True
            add_math_run_to_paragraph(p_q, q.get('cau_hoi', ''))

    list_tl = de.get("tu_luan", [])
    if list_tl:
        p_hdr4 = doc.add_paragraph()
        p_hdr4.add_run(f"\nPHẦN IV. Tự luận ({config.get('score_part4', 3.0)} điểm).").bold = True
        for q in list_tl:
            p_q = doc.add_paragraph()
            muc_do = q.get('muc_do', 'VDC')
            p_q.add_run(f"Câu {q.get('id', '')} ({muc_do}_TL): ").bold = True
            add_math_run_to_paragraph(p_q, q.get('cau_hoi', ''))

    doc.add_page_break()
    doc.add_paragraph().add_run(f"HƯỚNG DẪN CHẤM & ĐÁP ÁN - MÃ ĐỀ: {code_label}\n").bold = True
    da = data.get('dap_an_chi_tiet', {})
    
    if da.get('trac_nghiem_4_lua_chon'):
        doc.add_paragraph().add_run("Đáp án Phần I:").bold = True
        ans_table1 = doc.add_table(rows=1, cols=2)
        ans_table1.style = 'Table Grid'
        ans_table1.rows[0].cells[0].text = "Câu hỏi"
        ans_table1.rows[0].cells[1].text = "Đáp án"
        for q_id, val in da['trac_nghiem_4_lua_chon'].items():
            rc = ans_table1.add_row().cells
            rc[0].text = f"Câu {q_id}"
            rc[1].text = str(val)

    if da.get('tu_luan'):
        doc.add_paragraph().add_run("\nHướng dẫn giải chi tiết Phần IV (Tự luận):").bold = True
        for q_id, detail in da['tu_luan'].items() if isinstance(da['tu_luan'], dict) else enumerate(da['tu_luan']):
            p_tl = doc.add_paragraph()
            p_tl.add_run(f"Câu {q_id}: ").bold = True
            add_math_run_to_paragraph(p_tl, str(detail))
            
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ==========================================
# GIAO DIỆN ĐIỀU KHIỂN STREAMLIT
# ==========================================
st.markdown('<div class="main-title">Trợ Lý Ra Đề Kiểm Tra</div>', unsafe_allow_html=True)
st.info('Hỗ trợ GV ra đề kiểm tra kèm Ma trận và Bảng đặc tả')
if "gemini_api_key" in st.session_state and st.session_state["gemini_api_key"].strip() != "":
    api_key_input = st.session_state["gemini_api_key"]
else:
    st.warning("⚠️ Vui lòng cấu hình Google Gemini API Key tại Trang chủ trước khi vận hành.")
    st.stop()

model = init_gemini_client(api_key_input)

if 'current_document_content' not in st.session_state: 
    st.session_state.current_document_content = ""

tab1, tab2, tab3 = st.tabs(["| 📋 1. Chọn Môn học & Số câu", "| 📊 2. Phân bổ Ma trận & Điểm số", "| 📥 3. Đóng gói & Xuất Đề"])

with tab1:
    col1, col2 = st.columns(2)
    with col1:
        with st.container(border=True):
            st.markdown(
        """
        <div style="background-color: #E0F2FE; padding: 4px; border-left: 5px solid #0284C7; border-radius: 4px; margin-bottom: 10px;">
            <h4 style="margin: 0; color: #0369A1;">Lựa chọn bộ môn & Thông tin chung</h4>
        </div>
        """, 
        unsafe_allow_html=True
    )
            subject = st.selectbox("Chọn môn học cần thiết lập đề thi:", list(SUBJECTS_CONFIG.keys()))
            grade = st.selectbox("Khối lớp học:", [str(i) for i in range(1, 13)], index=7)
            exam_type = st.selectbox("Hình thức kiểm tra:", ["15 phút", "45 phút", "Giữa học kỳ I", "Cuối học kỳ I", "Giữa học kỳ II", "Cuối học kỳ II"])
            duration = st.number_input("Thời lượng làm bài (phút):", min_value=15, max_value=150, value=15, step=5)
            school_year = st.text_input("Năm học:", value="2026-2027")
            
            difficulty_levels = {
                "Cơ bản (Phù hợp HS Yếu - Trung bình)": {"nb": 50, "th": 30, "vd": 20, "vdc": 0},
                "Tiêu chuẩn (Bám sát SGK, đại trà)": {"nb": 40, "th": 30, "vd": 20, "vdc": 10},
                "Nâng cao (Phù hợp lớp khá - lớp chọn)": {"nb": 30, "th": 30, "vd": 25, "vdc": 15},
                "Phân hóa mạnh (Bồi dưỡng giỏi)": {"nb": 20, "th": 30, "vd": 30, "vdc": 20}
            }
            level_choice = st.selectbox("Độ khó mục tiêu của đề thi:", list(difficulty_levels.keys()), index=1)
            target_ratios = difficulty_levels[level_choice]

    with col2:
        with st.container(border=True):
            st.markdown(
        """
        <div style="background-color: #E0F2FE; padding: 4px; border-left: 5px solid #0284C7; border-radius: 4px; margin-bottom: 10px;">
            <h4 style="margin: 0; color: #0369A1;">Cấu hình số lượng câu hỏi</h4>
        </div>
        """, 
        unsafe_allow_html=True
    )
            
            # Lựa chọn cấu trúc đề thi chính
            exam_format = st.selectbox(
                "Chọn cấu trúc / dạng đề kiểm tra:",
                ["Trắc nghiệm + Tự luận", "Trắc nghiệm 100%", "Tự luận 100%"]
            )
            
            # Xử lý logic số lượng câu hỏi theo dạng đề
            if exam_format == "Trắc nghiệm 100%":
                num_tn_4_lua_chon = st.number_input("Trắc nghiệm nhiều lựa chọn (Phần I):", min_value=0, max_value=40, value=12)
                num_tn_dung_sai = st.number_input("Trắc nghiệm Đúng/Sai (Phần II):", min_value=0, max_value=10, value=2, help="Mỗi câu dạng này có 4 ý Đúng / Sai")
                num_tn_tra_loi_ngan = st.number_input("Trắc nghiệm Trả lời ngắn (Phần III):", min_value=0, max_value=15, value=4)
                num_tl = 0
                st.caption("ℹ️ *Đã khóa và đặt số câu Tự luận về 0.*")
                
            elif exam_format == "Tự luận 100%":
                num_tn_4_lua_chon = 0
                num_tn_dung_sai = 0
                num_tn_tra_loi_ngan = 0
                num_tl = st.number_input("Số câu hỏi Tự luận (Phần IV):", min_value=1, max_value=15, value=4)
                st.caption("ℹ️ *Đã khóa và đặt số câu Trắc nghiệm về 0.*")
                
            else:
                num_tn_4_lua_chon = st.number_input("Trắc nghiệm nhiều lựa chọn (Phần I):", min_value=0, max_value=40, value=12)
                num_tn_dung_sai = st.number_input("Trắc nghiệm Đúng/Sai (Phần II):", min_value=0, max_value=10, value=2, help="Mỗi câu dạng này có 4 ý Đúng / Sai")
                num_tn_tra_loi_ngan = st.number_input("Trắc nghiệm Trả lời ngắn (Phần III):", min_value=0, max_value=15, value=4)
                num_tl = st.number_input("Số câu hỏi Tự luận (Phần IV):", min_value=0, max_value=10, value=3)
            
            st.markdown('---')
            code_choice = st.selectbox("Số lượng mã đề đảo tự động:", [1, 2, 4, 6, 8], index=2)
            num_codes = int(code_choice)
            code_prefix = st.text_input("Ký hiệu mã đề bắt đầu:", value="101")

with tab2:
    st.markdown(
        """
        <div style="background-color: #E0F2FE; padding: 4px; border-left: 5px solid #0284C7; border-radius: 4px; margin-bottom: 10px;">
            <h4 style="margin: 0; color: #0369A1;">Nguồn nội dung kiến thức ra đề</h4>
        </div>
        """, 
        unsafe_allow_html=True
    )
    st.markdown('<div class="section-header"></div>', unsafe_allow_html=True)
    content_source = st.radio("Chọn phương thức cung cấp nội dung:", ["Nhập tay danh sách chủ đề", "Upload file tài liệu đa phương thức"], horizontal=True)
    
    if content_source == "Nhập tay danh sách chủ đề":
        topics_list = st.text_area("Danh sách các chủ đề kiến thức:", value="Chương 1: Khái niệm cơ bản\nChương 2: Bài toán vận dụng liên quan", height=100)
        st.session_state.current_document_content = topics_list
    else:
        uploaded_doc = st.file_uploader("Tải lên tài liệu:", type=["docx", "txt", "pdf", "png", "jpg", "jpeg"])
        user_custom_req = st.text_input("Yêu cầu bổ sung khi đọc tài liệu:", value="")
        st.session_state.user_custom_req = user_custom_req.strip()
        
        if uploaded_doc is not None:
            file_bytes = uploaded_doc.read()
            file_ext = uploaded_doc.name.split('.')[-1].lower()
            if file_ext == "txt": st.session_state.current_document_content = str(file_bytes, "utf-8")
            elif file_ext == "docx": st.session_state.current_document_content = extract_text_from_docx(file_bytes)
            elif file_ext == "pdf": st.session_state.current_document_content = {"mime_type": "application/pdf", "data": file_bytes}
            elif file_ext in ["png", "jpg", "jpeg"]: st.session_state.current_document_content = {"mime_type": "image/png" if file_ext == "png" else "image/jpeg", "data": file_bytes}

    # ==========================================
    # ĐỒNG BỘ LOGIC PHÂN BỔ ĐIỂM SỐ TỰ ĐỘNG
    # ==========================================
    st.markdown(
        """
        <div style="background-color: #E0F2FE; padding: 4px; border-left: 5px solid #0284C7; border-radius: 4px; margin-bottom: 10px;">
            <h6 style="margin: 0; color: #0369A1;">Phân bổ điểm số cho từng phần & Tùy chọn điểm Vận dụng cao</h6>
        </div>
        """, 
        unsafe_allow_html=True
    )
    st.markdown('<div class="section-header"></div>', unsafe_allow_html=True)
    
    col_s1, col_s2, col_s3, col_s4, col_vdc = st.columns(5)
    
    if exam_format == "Trắc nghiệm 100%":
        with col_s1: score_part1 = st.number_input("Điểm Phần I (Trắc nghiệm):", min_value=0.0, max_value=10.0, value=4.0, step=0.25)
        with col_s2: score_part2 = st.number_input("Điểm Phần II (Đúng/Sai):", min_value=0.0, max_value=10.0, value=3.0, step=0.25)
        with col_s3: score_part3 = st.number_input("Điểm Phần III (Trả lời ngắn):", min_value=0.0, max_value=10.0, value=3.0, step=0.25)
        score_part4 = 0.0  # Khóa cứng Tự luận
        with col_s4: st.number_input("Điểm Phần IV (Tự luận):", value=0.0, disabled=True)
        with col_vdc: score_vdc_custom = st.number_input("Điểm dành riêng cho VDC:", min_value=0.0, max_value=5.0, value=1.0, step=0.25)

    elif exam_format == "Tự luận 100%":
        score_part1 = 0.0
        score_part2 = 0.0
        score_part3 = 0.0
        with col_s1: st.number_input("Điểm Phần I (Trắc nghiệm):", value=0.0, disabled=True)
        with col_s2: st.number_input("Điểm Phần II (Đúng/Sai):", value=0.0, disabled=True)
        with col_s3: st.number_input("Điểm Phần III (Trả lời ngắn):", value=0.0, disabled=True)
        with col_s4: score_part4 = st.number_input("Điểm Phần IV (Tự luận):", min_value=10.0, max_value=10.0, value=10.0, disabled=True, help="Thuần tự luận mặc định là 10 điểm")
        with col_vdc: score_vdc_custom = st.number_input("Điểm dành riêng cho VDC:", min_value=0.0, max_value=5.0, value=2.0, step=0.25)

    else: # Trắc nghiệm + Tự luận kết hợp
        with col_s1: score_part1 = st.number_input("Điểm Phần I (Trắc nghiệm):", min_value=0.0, max_value=10.0, value=3.0, step=0.25)
        with col_s2: score_part2 = st.number_input("Điểm Phần II (Đúng/Sai):", min_value=0.0, max_value=10.0, value=2.0, step=0.25)
        with col_s3: score_part3 = st.number_input("Điểm Phần III (Trả lời ngắn):", min_value=0.0, max_value=10.0, value=2.0, step=0.25)
        with col_s4: score_part4 = st.number_input("Điểm Phần IV (Tự luận):", min_value=0.0, max_value=10.0, value=3.0, step=0.25)
        with col_vdc: score_vdc_custom = st.number_input("Điểm dành riêng cho VDC:", min_value=0.0, max_value=5.0, value=1.0, step=0.25)
    
    total_score = score_part1 + score_part2 + score_part3 + score_part4
    
    if total_score > 10.0:
        st.error(f"❌ LỖI CẤU HÌNH ĐIỂM: Tổng số điểm hiện tại là **{total_score}** điểm, vượt quá giới hạn cho phép (Tối đa 10 điểm). Vui lòng điều chỉnh lại!")
        score_error = True
    elif total_score < 10.0:
        st.warning(f"⚠️ Tổng số điểm hiện tại là **{total_score}** điểm. Đề kiểm tra chuẩn cần đạt chính xác **10.0** điểm.")
        score_error = False
    else:
        st.success("✅ Cấu hình điểm số chính xác đạt 10/10 điểm chuẩn.")
        score_error = False

    st.markdown(
        """
        <div style="background-color: #E0F2FE; padding: 4px; border-left: 5px solid #0284C7; border-radius: 4px; margin-bottom: 10px;">
            <h4 style="margin: 0; color: #0369A1;">Lựa chọn mẫu cấu trúc Ma trận hiển thị</h4>
        </div>
        """, 
        unsafe_allow_html=True
    )
    matrix_template = st.radio("Chọn mẫu ma trận xuất bản:", ["Mẫu đơn giản truyền thống", "Mẫu quy chuẩn Công văn 7991/BGDĐT-GDTrH"], index=0, horizontal=True)

    st.markdown(
        """
        <div style="background-color: #E0F2FE; padding: 4px; border-left: 5px solid #0284C7; border-radius: 4px; margin-bottom: 10px;">
            <h4 style="margin: 0; color: #0369A1;">Phân bổ Tỷ lệ Ma trận tư duy</h4>
        </div>
        """, 
        unsafe_allow_html=True
    )
    c1, c2, c3, c4 = st.columns(4)
    with c1: nb_ratio = st.slider("Nhận biết (%)", 0, 100, target_ratios["nb"])
    with c2: th_ratio = st.slider("Thông hiểu (%)", 0, 100, target_ratios["th"])
    with c3: vd_ratio = st.slider("Vận dụng (%)", 0, 100, target_ratios["vd"])
    with c4: vdc_ratio = st.slider("Vận dụng cao (%)", 0, 100, target_ratios["vdc"])
    
    total_ratio = nb_ratio + th_ratio + vd_ratio + vdc_ratio
    if total_ratio != 100:
        st.warning(f"⚠️ Tổng tỷ lệ hiện tại đạt {total_ratio}%. Thầy cô vui lòng căn chỉnh về chính xác 100%.")

    st.markdown("---")
    col_btn1, col_btn2 = st.columns(2)
    
    with col_btn1:
        if st.button("📊 BƯỚC 1: KHỞI TẠO MA TRẬN & ĐẶC TẢ", type="primary", use_container_width=True, disabled=score_error):
            if total_ratio != 100:
                st.error("Tổng tỷ lệ phần trăm phân bổ điểm phải bằng 100% trước khi khởi tạo.")
            elif not st.session_state.current_document_content:
                st.error("Vui lòng cung cấp danh sách chủ đề hoặc tải file tài liệu đính kèm lên trước.")
            else:
                config_pkg = {
                    "subject": subject, "grade": grade, "duration": duration, "num_tn_4_lua_chon": num_tn_4_lua_chon, 
                    "num_tn_dung_sai": num_tn_dung_sai, "num_tn_tra_loi_ngan": num_tn_tra_loi_ngan, "num_tl": num_tl,
                    "score_part1": score_part1, "score_part2": score_part2, "score_part3": score_part3, "score_part4": score_part4,
                    "score_vdc_custom": score_vdc_custom,
                    "nb_ratio": nb_ratio, "th_ratio": th_ratio, "vd_ratio": vd_ratio, "vdc_ratio": vdc_ratio,
                    "difficulty": level_choice, "matrix_template": matrix_template
                }
                with st.spinner("AI đang tính toán khung ma trận đặc tả..."):
                    try:
                        st.session_state.step1_data = generate_step1_matrix(model, config_pkg, st.session_state.current_document_content)
                        st.success("✅ Đã thiết lập xong Khung đặc tả bộ môn!")
                    except Exception as e:
                        st.error(f"Lỗi phân tích kết nối: {e}")

    if st.session_state.step1_data:
        with col_btn2:
            if st.button("🔥 BƯỚC 2: SINH CÂU HỎI & ĐẢO MÃ ĐỀ", disabled=score_error):
                config_pkg = {
                    "subject": subject, "grade": grade, "duration": duration, "num_tn_4_lua_chon": num_tn_4_lua_chon, 
                    "num_tn_dung_sai": num_tn_dung_sai, "num_tn_tra_loi_ngan": num_tn_tra_loi_ngan, "num_tl": num_tl,
                    "score_part1": score_part1, "score_part2": score_part2, "score_part3": score_part3, "score_part4": score_part4,
                    "score_vdc_custom": score_vdc_custom,
                    "nb_ratio": nb_ratio, "th_ratio": th_ratio, "vd_ratio": vd_ratio, "vdc_ratio": vdc_ratio,
                    "difficulty": level_choice, "matrix_template": matrix_template
                }
                with st.spinner("AI đang bám sát Khung đặc tả để soạn nội dung câu hỏi kiểm tra..."):
                    try:
                        rule = SUBJECTS_CONFIG[subject]
                        step2_data = generate_step2_questions(model, config_pkg, st.session_state.step1_data, rule, st.session_state.current_document_content)
                        
                        full_data = {
                            "ma_tran": st.session_state.step1_data["ma_tran"],
                            "bang_dac_ta": st.session_state.step1_data["bang_dac_ta"],
                            "de_kiem_tra": step2_data["de_kiem_tra"],
                            "dap_an_chi_tiet": step2_data["dap_an_chi_tiet"]
                        }
                        
                        try: s_code = int(code_prefix)
                        except ValueError: s_code = 101
                        
                        bundle, alignment_df = generate_shuffled_bundle(full_data, s_code, num_codes)
                        bundle["ĐỀ GỐC"] = full_data
                        st.session_state.multi_codes_data = bundle
                        st.session_state.alignment_table = alignment_df
                        st.success(f"🎉 Đã tạo xong ĐỀ GỐC và {num_codes} mã đề đảo.")
                    except Exception as e:
                        st.error(f"Lỗi tạo câu hỏi: {e}")
                        
with tab3:
    if not st.session_state.get('multi_codes_data'):
        st.info("Hệ thống đang ở trạng thái chờ nạp dữ liệu câu hỏi từ Bước 2.")
    else:
        config_pkg = { 
            "subject": subject, "grade": grade, "exam_type": exam_type, "duration": duration, "school_year": school_year,
            "score_part1": score_part1, "score_part2": score_part2, "score_part3": score_part3, "score_part4": score_part4,
            "matrix_template": matrix_template, "nb_ratio": nb_ratio, "th_ratio": th_ratio, "vd_ratio": vd_ratio, "vdc_ratio": vdc_ratio
        }
        inc_mat = st.checkbox("Chèn bảng Ma trận phân bổ vào đầu tài liệu", value=True)
        export_mode = st.radio("Hình thức xuất bản bộ đề:", ["Tải file đơn lẻ", "Nén tất cả mã đề vào file ZIP", "Gộp chung vào 1 file Word"])
        
        if export_mode == "Tải file đơn lẻ":
            all_available_codes = list(st.session_state.multi_codes_data.keys())
            if "ĐỀ GỐC" in all_available_codes:
                all_available_codes.remove("ĐỀ GỐC")
                all_available_codes = ["ĐỀ GỐC"] + all_available_codes
            sel_code = st.selectbox("Chọn bản đề cần xuất bản:", all_available_codes)
            docx_buf = build_single_docx(config_pkg, st.session_state.multi_codes_data[sel_code], sel_code, include_matrix=inc_mat)
            st.download_button(label=f"📥 TẢI FILE WORD [{sel_code}]", data=docx_buf, file_name=f"De_{subject}_{sel_code}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
        elif export_mode == "Nén tất cả mã đề vào file ZIP":
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
                for c_code, c_data in st.session_state.multi_codes_data.items():
                    d_buf = build_single_docx(config_pkg, c_data, c_code, include_matrix=inc_mat)
                    filename_in_zip = f"De_Goc_{subject}_Lop{grade}.docx" if c_code == "ĐỀ GỐC" else f"De_{subject}_Lop{grade}_MaDe_{c_code}.docx"
                    zip_file.writestr(filename_in_zip, d_buf.getvalue())
            zip_buffer.seek(0)
            st.download_button(label="📥 TẢI FILE NÉN ZIP TRỌN BỘ", data=zip_buffer, file_name=f"Bo_De_{subject}_TronGoi.zip", mime="application/zip")
            
        else:
            main_doc = Document()
            all_codes = list(st.session_state.multi_codes_data.keys())
            if "ĐỀ GỐC" in all_codes:
                all_codes.remove("ĐỀ GỐC")
                all_codes = ["ĐỀ GỐC"] + all_codes
            for idx, c_code in enumerate(all_codes):
                c_data = st.session_state.multi_codes_data[c_code]
                temp_buf = build_single_docx(config_pkg, c_data, c_code, include_matrix=inc_mat)
                t_doc = Document(temp_buf)
                for element in t_doc.element.body: main_doc.element.body.append(element)
                if idx < len(all_codes) - 1: main_doc.add_page_break()
                
            all_buf = BytesIO()
            main_doc.save(all_buf)
            all_buf.seek(0)
            st.download_button(label="📥 TẢI FILE GỘP TOÀN BỘ (.DOCX)", data=all_buf, file_name=f"Gop_Bo_De_{subject}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# --- FOOTER CỐ ĐỊNH ---
st.divider()
st.markdown("---")

col_left, col_right = st.columns(2)
with col_left:
    st.caption("Phát triển bởi Ngo Thanh Hung © 2026")
with col_right:
    st.markdown(
        "<div style='text-align: right; color: gray; font-size: 0.85em;'>"
        "AI có thể mắc lỗi. Cần kiểm tra kỹ các thông tin quan trọng."
        "</div>", 
        unsafe_allow_html=True
    )
