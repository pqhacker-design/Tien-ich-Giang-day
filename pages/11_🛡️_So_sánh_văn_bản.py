import streamlit as st
import os
import io
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from ai_auditor_app.core.document_proc import DocumentProcessor
from ai_auditor_app.core.ai_engine import AIEngine

load_dotenv()

st.set_page_config(page_title="AI So sánh - Đối chiếu Văn bản Giáo dục", layout="wide")

# Khởi tạo Session State lưu trữ dữ liệu liên tục và đồng bộ
if "template_reqs" not in st.session_state:
    st.session_state.template_reqs = []
if "audit_report" not in st.session_state:
    st.session_state.audit_report = None
if "user_text_content" not in st.session_state:
    st.session_state.user_text_content = ""
if "accepted_fixes" not in st.session_state:
    st.session_state.accepted_fixes = []

st.markdown("## 🛡️ AI So sánh văn bản với mẫu hướng dẫn")
st.caption("Ứng dụng tự động đối chiếu, phát hiện mục thiếu sót và hỗ trợ hiệu đính hồ sơ theo tiêu chuẩn.")

# --- SIDEBAR: Cấu hình hệ thống ---
if "gemini_api_key" in st.session_state and st.session_state["gemini_api_key"].strip() != "":
    api_key_input = st.session_state["gemini_api_key"].strip()
    os.environ["GEMINI_API_KEY"] = api_key_input
    
    st.sidebar.header("⚙️ Cấu hình Mô hình")
    model_name = st.sidebar.selectbox("Lựa chọn Mô hình", ["gemini-2.5-flash", "gemini-3.0-flash"])
    audit_level = st.sidebar.radio("Mức độ rà soát", ["Toàn diện (Cấu trúc + Câu từ)", "Cấu trúc khung", "Từ khóa & Minh chứng"])
    
    if 'ai_engine' not in st.session_state or st.session_state.ai_engine is None:
        st.session_state.ai_engine = AIEngine(api_key=api_key_input, model_name=model_name)
    else:
        st.session_state.ai_engine.model_name = model_name
    
    ai_engine = st.session_state.ai_engine
else:
    st.warning("⚠️ Vui lòng quay lại **Trang chủ** để nhập Google Gemini API Key trước khi sử dụng tính năng này.")
    st.info("💡 Mẹo: Nhập một lần tại trang chủ, tất cả các công cụ khác sẽ tự động kích hoạt.")
    st.page_link("🏠_Trang_Chủ.py", label="Nhấn vào đây để Quay lại Trang chủ", icon="🔄")
    st.stop()

# --- HÀM XUẤT FILE BIÊN BẢN ĐÁNH GIÁ VĂN BẢN ---
def export_audit_to_docx(audit_report, avg_score=None):
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_header = p_header.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n" + "_"*15 + "\n")
    run_header.bold = True
    
    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("BIÊN BẢN THẨM ĐỊNH & ĐỐI CHIẾU HỒ SƠ GIÁO DỤC\n(Hệ thống AI tự động phân tích)")
    run_title.font.size = Pt(16)
    run_title.font.color.rgb = None
    
    if avg_score:
        doc.add_paragraph(f"Điểm đánh giá tổng hợp: {round(avg_score, 1)} / 100 điểm.")
    
    doc.add_paragraph("Chi tiết kết quả đối chiếu so với văn bản mẫu:").bold = True
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'STT'
    hdr_cells[1].text = 'Mục/Tiêu chuẩn đối chiếu'
    hdr_cells[2].text = 'Trạng thái Đánh giá'
    hdr_cells[3].text = 'Khuyến nghị bổ sung / Chỉnh sửa'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for idx, item in enumerate(audit_report):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx + 1)
        row_cells[1].text = str(item.get('criteria', ''))
        row_cells[2].text = f"{item.get('status', '')} ({item.get('score', 70)}/100)"
        row_cells[3].text = str(item.get('fix', ''))
        
    doc.add_paragraph("\n").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_footer.add_run("HỆ THỐNG TRỢ LÝ AI BIÊN TẬP\n(Tự động xuất văn bản)").italic = True
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- HÀM VÁ LỖI TRỰC TIẾP VÀO FILE NGƯỜI DÙNG ---
def export_fixed_doc(uploaded_file, audit_report):
    uploaded_file.seek(0)
    file_bytes = uploaded_file.read()
    if not file_bytes:
        return io.BytesIO()
        
    doc = Document(io.BytesIO(file_bytes))
    pending_fixes = [item for item in audit_report if item.get('status') != "Đạt"]
    
    if pending_fixes:
        for item in pending_fixes:
            criteria_name = item.get('criteria', '').lower().strip()
            fix_content = item.get('fix', '').strip()
            matched = False
            
            for paragraph in doc.paragraphs:
                if criteria_name in paragraph.text.lower() and len(paragraph.text) < 250:
                    p_fixed = doc.add_paragraph(style=paragraph.style)
                    paragraph._element.addnext(p_fixed._element)
                    
                    run_intro = p_fixed.add_run(f"\n[AI Sửa đổi/Bổ sung mục {item.get('criteria')}]: ")
                    run_intro.bold = True
                    run_intro.font.color.rgb = RGBColor(255, 0, 0)
                    
                    run_content = p_fixed.add_run(f"{fix_content}\n")
                    run_content.italic = True
                    run_content.font.color.rgb = RGBColor(255, 0, 0)
                    
                    matched = True
                    break
            
            if not matched:
                p_sec = doc.add_paragraph()
                r_title = p_sec.add_run(f"\n[AI Bổ sung mục còn thiếu: {item.get('criteria')}]\n")
                r_title.bold = True
                r_title.font.color.rgb = RGBColor(255, 0, 0)
                
                r_content = p_sec.add_run(f"{fix_content}\n")
                r_content.italic = True
                r_content.font.color.rgb = RGBColor(255, 0, 0)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- GIAO DIỆN TẬP TRUNG (3 TABS) ---
tab1, tab2, tab3 = st.tabs([
    "📋 1. Văn Bản Mẫu Chuẩn", 
    "📂 2. Hồ Sơ Người Dùng", 
    "🔬 3. Kết Quả Đối Chiếu & Xuất Đánh Giá"
])

# TAB 1: VĂN BẢN MẪU
with tab1:
    st.markdown(
        """
        <div style="background-color: #E0F2FE; padding: 4px; border-left: 5px solid #0284C7; border-radius: 4px; margin-bottom: 10px;">
            <h4 style="margin: 0; color: #0369A1;">📂 Cung cấp Văn bản mẫu chỉ dẫn / Tiêu chuẩn gốc</h4>
        </div>
        """, 
        unsafe_allow_html=True
    )
    uploaded_template = st.file_uploader("Tải lên File Mẫu Quy chuẩn (DOCX/PDF)", type=["docx", "pdf"], key="tmpl")
    if uploaded_template and ai_engine:
        with st.spinner("AI đang phân tích cấu trúc mẫu quy chuẩn..."):
            text = DocumentProcessor.read_file(uploaded_template)
            st.session_state.template_reqs = ai_engine.extract_template_requirements(text)
            st.success(f"Đã trích xuất thành công {len(st.session_state.template_reqs)} yêu cầu cốt lõi từ văn bản mẫu.")
            st.json(st.session_state.template_reqs)

# TAB 2: VĂN BẢN CẦN KIỂM TRA
with tab2:
    st.markdown(
        """
        <div style="background-color: #E0F2FE; padding: 4px; border-left: 5px solid #0284C7; border-radius: 4px; margin-bottom: 10px;">
            <h4 style="margin: 0; color: #0369A1;">📂 Tải lên hồ sơ/văn bản cần thẩm định đối chiếu</h4>
        </div>
        """, 
        unsafe_allow_html=True
    )
    uploaded_user_file = st.file_uploader("Tải lên Hồ sơ của bạn (DOCX/PDF)", type=["docx", "pdf"], key="user_doc")
    if uploaded_user_file:
        raw_text = DocumentProcessor.read_file(uploaded_user_file)
        if raw_text.startswith("LỖI:"):
            st.error(raw_text)
            st.session_state.user_text_content = ""
        else:
            st.session_state.user_text_content = raw_text
            st.success("Tải hồ sơ cá nhân thành công!")
            st.text_area("Xem trước Nội dung văn bản hiện tại:", value=st.session_state.user_text_content[:1500] + "\n...", height=300)

# TAB 3: PHÂN TÍCH, ĐỐI CHIẾU & XUẤT BÁO CÁO
with tab3:
    st.markdown(
        """
        <div style="background-color: #E0F2FE; padding: 4px; border-left: 5px solid #0284C7; border-radius: 4px; margin-bottom: 10px;">
            <h4 style="margin: 0; color: #0369A1;">📂 Kết quả phân tích, phát hiện sai sót & bổ sung</h4>
        </div>
        """, 
        unsafe_allow_html=True
    )
    if st.button("🚀 Khởi chạy Rà soát & So sánh diện rộng bằng AI"):
        if st.session_state.user_text_content and st.session_state.template_reqs and ai_engine:
            with st.spinner("AI đang tiến hành đối chiếu cấu trúc, kiểm tra các mục thiếu hụt..."):
                res = ai_engine.audit_document(st.session_state.template_reqs, st.session_state.user_text_content, audit_level)
                st.session_state.audit_report = res.get("report", [])
                st.session_state.accepted_fixes = [] # Reset lại danh sách duyệt sửa
        else:
            st.warning("Vui lòng đảm bảo đã nạp đầy đủ Văn bản mẫu (Tab 1) và Hồ sơ cần kiểm tra (Tab 2).")
            
    if st.session_state.audit_report:
        st.subheader("📊 Bảng đánh giá chi tiết sự phù hợp:")
        
        # Hiển thị trực quan dưới dạng biểu bảng dữ liệu Streamlit
        df_display = pd.DataFrame(st.session_state.audit_report)
        if not df_display.empty:
            st.dataframe(df_display[['criteria', 'status', 'score', 'fix']], use_container_width=True)
        
        st.write("---")
        st.subheader("🔍 Chi tiết từng tiêu chí và Tùy chọn Phê duyệt hiệu đính:")
        
        for idx, item in enumerate(st.session_state.audit_report):
            status_signal = "🟢 Đạt" if item.get('status') == "Đạt" else f"🔴 {item.get('status')}"
            with st.expander(f"📌 {item.get('criteria')} — Trạng thái: {status_signal}"):
                st.write(f"**Điểm độ tương thích:** {item.get('score')}/100")
                st.info(f"💡 **Hướng dẫn bổ sung / sửa đổi từ AI:** {item.get('fix')}")
                
                col_acc, col_rej = st.columns(2)
                if col_acc.button("☑️ Chấp nhận sửa mục này", key=f"acc_{idx}"):
                    if item not in st.session_state.accepted_fixes:
                        st.session_state.accepted_fixes.append(item)
                    st.success(f"Đã đưa mục '{item.get('criteria')}' vào danh sách vá lỗi trực tiếp.")
                    
                if col_rej.button("✖️ Giữ nguyên văn bản gốc", key=f"rej_{idx}"):
                    if item in st.session_state.accepted_fixes:
                        st.session_state.accepted_fixes.remove(item)
                    st.toast("Đã bỏ qua gợi ý sửa đổi.")

        st.divider()
        st.subheader("💾 Trung tâm Xuất & Tải File Hoàn Thiện")
        
        col_download_1, col_download_2 = st.columns(2)
        with col_download_1:
            avg_score = sum([int(i.get('score', 70)) for i in st.session_state.audit_report]) / len(st.session_state.audit_report)
            docx_report_data = export_audit_to_docx(st.session_state.audit_report, avg_score)
            st.download_button(
                label="📥 Tải Biên Bản Đánh Giá So Sánh (.DOCX)",
                data=docx_report_data,
                file_name="Bien_Ban_Danh_Gia_Doi_Chieu_AI.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="btn_download_report"
            )
            
        with col_download_2:
            if uploaded_user_file:
                with st.spinner("Hệ thống đang chèn các nội dung chỉnh sửa/bổ sung vào file gốc..."):
                    reports_to_fix = st.session_state.accepted_fixes if st.session_state.accepted_fixes else st.session_state.audit_report
                    fixed_docx_data = export_fixed_doc(uploaded_user_file, reports_to_fix)
                    
                st.download_button(
                    label="✨ Tải Văn Bản Đã Tự Động Vá Lỗi Mới (.DOCX)",
                    data=fixed_docx_data,
                    file_name="Van_Ban_User_Da_Sua_Doi.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="btn_download_fixed_doc"
                )

# --- FOOTER CỐ ĐỊNH ---
st.divider()
col_left, col_right = st.columns(2)
with col_left:
    st.caption("Phát triển bởi Ngo Thanh Hung © 2026")
with col_right:
    st.markdown(
        "<div style='text-align: right; color: gray; font-size: 0.85em;'>"
        "Hệ thống so sánh dựa trên thuật toán LLM. Cần kiểm tra kỹ trước khi ban hành văn bản."
        "</div>", 
        unsafe_allow_html=True
    )
