import streamlit as st
import io
import os
import pandas as pd
from docx import Document

# --- MODULE 1: XỬ LÝ WORD GIỮ NGUYÊN ĐỊNH DẠNG (Bê từ bản Desktop qua) ---
class WordProcessor:
    @staticmethod
    def extract_full_text(doc):
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        return "\n".join(full_text)

    @staticmethod
    def replace_text_keep_format(doc, search_text, replace_text):
        # Sửa đổi cấp Paragraph và Run cho văn bản thường
        for paragraph in doc.paragraphs:
            if search_text in paragraph.text:
                for run in paragraph.runs:
                    if search_text in run.text:
                        run.text = run.text.replace(search_text, replace_text)
        # Sửa đổi bên trong các bảng biểu
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if search_text in paragraph.text:
                            for run in paragraph.runs:
                                if search_text in run.text:
                                    run.text = run.text.replace(search_text, replace_text)
        return doc

# --- MODULE 2: KẾT NỐI GEMINI AI XỬ LÝ NGỮ CẢNH ---
class AICorrector:
    def __init__(self, api_key):
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-2.5-flash', generation_config={"response_mime_type": "application/json"})

    def analyze_text(self, text):
        import json
        if not text.strip(): return []
        prompt = f"""
        Bạn là chuyên gia ngôn ngữ Tiếng Việt hành chính và giáo dục. Hãy rà soát đoạn văn bản sau để tìm lỗi chính tả, ngữ pháp, từ ngữ rườm rà:
        \"\"\"{text}\"\"\"
        Trả về kết quả duy nhất ở cấu trúc JSON:
        {{ "errors": [ {{ "sai": "từ lỗi", "dung": "sửa đúng", "loai": "Chính tả" hoặc "Ngữ pháp" hoặc "Diễn đạt", "ly_do": "giải thích" }} ] }}
        """
        try:
            response = self.model.generate_content(prompt)
            return json.loads(response.text).get("errors", [])
        except:
            return []

# --- GIAO DIỆN STREAMLIT HIỆN ĐẠI ---
st.title("📝 Trợ lý Sửa lỗi Chính tả & Diễn đạt Hành chính")
st.write("Quét lỗi ngữ cảnh sâu bằng AI và tự động sửa đổi trực tiếp trên các khối định dạng Word gốc.")

# Lấy API Key tập trung từ Trang chủ giống các ứng dụng khác
if "gemini_api_key" in st.session_state and st.session_state["gemini_api_key"].strip() != "":
    api_key = st.session_state["gemini_api_key"]
else:
    st.warning("⚠️ Vui lòng quay lại **Trang chủ** để nhập Google Gemini API Key trước khi sử dụng.")
    st.stop()

col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("📂 1. Tải tài liệu cần quét")
    uploaded_file = st.file_uploader("Kéo thả file Word .docx vào đây:", type=["docx"])
    
    if uploaded_file is not None:
        file_bytes = io.BytesIO(uploaded_file.read())
        st.success(f"✔️ Đã nhận file: {uploaded_file.name}")
        
        # LƯU TÊN FILE VÀO SESSION STATE ĐỂ DÙNG CHUNG AN TOÀN
        st.session_state["current_file_name"] = uploaded_file.name
        
        if st.button("🚀 Bắt đầu quét lỗi văn bản", type="primary", use_container_width=True):
            with st.spinner("AI đang rà soát từng câu chữ..."):
                doc = Document(file_bytes)
                full_text = WordProcessor.extract_full_text(doc)
                
                # Gọi AI xử lý
                corrector = AICorrector(api_key)
                text_blocks = [full_text[i:i+4000] for i in range(0, len(full_text), 4000)]
                all_errors = []
                for block in text_blocks:
                    all_errors.extend(corrector.analyze_text(block))
                
                # Lưu trạng thái kết quả vào bộ nhớ Streamlit
                st.session_state["word_doc"] = doc
                st.session_state["word_errors"] = all_errors
                st.success(f"Quét xong! Tìm thấy {len(all_errors)} điểm cần tối ưu.")
    else:
        # KHI XÓA FILE: Tự động dọn dẹp các dữ liệu cũ trong bộ nhớ để tránh xung đột giao diện
        if "word_errors" in st.session_state:
            del st.session_state["word_errors"]
        if "word_doc" in st.session_state:
            del st.session_state["word_doc"]
        if "fixed_file" in st.session_state:
            del st.session_state["fixed_file"]
        if "current_file_name" in st.session_state:
            del st.session_state["current_file_name"]

with col2:
    st.subheader("📊 2. Kết quả phân tích & Sửa đổi")
    if "word_errors" in st.session_state and st.session_state["word_errors"]:
        errors = st.session_state["word_errors"]
        doc = st.session_state["word_doc"]
        
        # Lấy tên file an toàn từ session_state (mặc định là 'Tai_Lieu.docx' nếu trống)
        safe_file_name = st.session_state.get("current_file_name", "Tai_Lieu.docx")
        
        # Hiển thị bảng lưới lỗi trực quan
        df_show = pd.DataFrame(errors)
        df_show.columns = ["Từ bị sai / Kém hiệu quả", "Đề xuất sửa đúng", "Phân loại lỗi", "Lý do chi tiết"]
        st.dataframe(df_show, use_container_width=True)
        
        st.markdown("---")
        if st.button("🔄 Áp dụng sửa tất cả lỗi & Giữ nguyên định dạng gốc", type="secondary", use_container_width=True):
            with st.spinner("Hệ thống đang thay thế thông minh các khối văn bản..."):
                for err in errors:
                    doc = WordProcessor.replace_text_keep_format(doc, err["sai"], err["dung"])
                
                output_stream = io.BytesIO()
                doc.save(output_stream)
                st.session_state["fixed_file"] = output_stream.getvalue()
                st.toast("🎉 Đã sửa xong toàn bộ văn bản trên hệ thống!")

        if "fixed_file" in st.session_state:
            st.write("")
            st.download_button(
                label="💾 TẢI FILE WORD ĐÃ SỬA SẠCH LỖI (.DOCX)",
                data=st.session_state["fixed_file"],
                file_name=f"Sua_Loi_Chinh_Ta_{safe_file_name}", # <-- ĐÃ THAY BẰNG BIẾN AN TOÀN
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    else:
        st.info("Chưa có dữ liệu. Vui lòng hoàn thành bước nạp và quét file ở cột bên trái.")
