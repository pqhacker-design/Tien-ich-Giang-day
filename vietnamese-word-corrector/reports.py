from docx import Document
import pandas as pd

class ReportGenerator:
    @staticmethod
    def export_excel(error_list, output_path):
        """Xuất file thống kê Excel chi tiết"""
        data = []
        for idx, err in enumerate(error_list):
            data.append({
                "STT": idx + 1,
                "Nội dung lỗi": err["sai"],
                "Đề xuất sửa đúng": err["dung"],
                "Phân loại lỗi": err["loai"],
                "Lý do sai / Diễn đạt": err["ly_do"],
                "Trạng thái": err.get("status", "Chờ duyệt")
            })
        df = pd.DataFrame(data)
        df.to_excel(output_path, index=False, engine='openpyxl')

    @staticmethod
    def export_word_report(error_list, output_path):
        """Xuất biên bản báo cáo lỗi định dạng Word"""
        doc = Document()
        doc.add_heading("BÁO CÁO THỐNG KÊ LỖI CHÍNH TẢ & NGỮ PHÁP", level=1)
        doc.add_paragraph(f"Tổng số lỗi phát hiện: {len(error_list)}")
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nội dung lỗi'
        hdr_cells[1].text = 'Nội dung đúng'
        hdr_cells[2].text = 'Loại lỗi'
        hdr_cells[3].text = 'Lý do'
        
        for err in error_list:
            row_cells = table.add_row().cells
            row_cells[0].text = err["sai"]
            row_cells[1].text = err["dung"]
            row_cells[2].text = err["loai"]
            row_cells[3].text = err["ly_do"]
            
        doc.save(output_path)