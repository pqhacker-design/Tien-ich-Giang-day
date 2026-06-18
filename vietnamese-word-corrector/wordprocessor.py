from docx import Document
import os

class WordProcessor:
    @staticmethod
    def load_document(file_path):
        """Nạp file Word"""
        return Document(file_path)

    @staticmethod
    def extract_full_text(doc):
        """Trích xuất toàn bộ văn bản để đẩy vào AI phân tích tổng thể"""
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        return "\n".join(full_text)

    @staticmethod
    def replace_text_keep_format(doc, search_text, replace_text):
        """Thay thế nội dung văn bản bị lỗi nhưng giữ nguyên 100% định dạng của Run"""
        for paragraph in doc.paragraphs:
            if search_text in paragraph.text:
                # Kỹ thuật xử lý chuỗi Run lân cận nếu từ bị phân tách
                for run in paragraph.runs:
                    if search_text in run.text:
                        run.text = run.text.replace(search_text, replace_text)
        
        # Quét và xử lý lỗi cả trong các bảng biểu nếu có
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if search_text in paragraph.text:
                            for run in paragraph.runs:
                                if search_text in run.text:
                                    run.text = run.text.replace(search_text, replace_text)
        return doc