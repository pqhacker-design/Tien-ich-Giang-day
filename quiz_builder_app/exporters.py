import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import openpyxl

def export_to_word(exam_title, exam_structure, shuffled_questions):
    doc = docx.Document()
    
    # Thiết lập lề trang chuẩn kỹ thuật văn bản ban hành bộ GD
    for section in doc.sections:
        section.top_margin = Inches(0.79) # 2cm
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18) # 3cm
        section.right_margin = Inches(0.59) # 1.5cm

    # Tiêu đề Header Trường / Sở
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "SỞ GIÁO DỤC VÀ ĐÀO TẠO\nTRƯỜNG THPT CHUYÊN CHUẨN ĐỀ"
    hdr_cells[1].text = f"ĐỀ KIỂM TRA CHÍNH THỨC\nMôn: {exam_title}\nThời gian: 90 phút"
    
    doc.add_paragraph("\n")
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(f"MÃ ĐỀ THI KHẢO SÁT")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Họ và tên thí sinh: ....................................................... Số báo danh: .......................\n")
    
    # Ghi nội dung câu hỏi
    for idx, q in enumerate(shuffled_questions, 1):
        p = doc.add_paragraph()
        p.add_run(f"Câu {idx} ({q['level']}): ").bold = True
        p.add_run(q['content'])
        
        if q['shuffled_options']:
            for opt in q['shuffled_options']:
                doc.add_paragraph(f"   {opt}")
                
    # Lưu và trả về buffer
    target_stream = BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream

def export_matrix_to_excel(matrix_df):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Ma trận đề thi"
    
    # Ghi Header cột
    columns = list(matrix_df.columns)
    ws.append(columns)
    
    # Ghi dữ liệu dòng
    for _, row in matrix_df.iterrows():
        ws.append(list(row))
        
    excel_stream = BytesIO()
    wb.save(excel_stream)
    excel_stream.seek(0)
    return excel_stream
