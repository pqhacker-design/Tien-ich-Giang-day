import streamlit as st
import io
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Khối import linh hoạt để tránh lỗi đường dẫn hệ thống trên Streamlit Cloud
try:
    from ai_engine import call_ai_stream
except ModuleNotFoundError:
    from edu_research_assistant.ai_engine import call_ai_stream

def export_refined_docx(topic, refined_text):
    """Đóng gói văn bản đã được AI hiệu đính sang file Word chuẩn Nghị định 30/2020/NĐ-CP"""
    doc = Document()
    
    # 1. Cấu hình lề trang chuẩn hành chính (Top 20mm, Bottom 20mm, Left 30mm, Right 15mm)
    for section in doc.sections:
        section.top_margin = Inches(20 / 25.4)
        section.bottom_margin = Inches(20 / 25.4)
        section.left_margin = Inches(30 / 25.4)
        section.right_margin = Inches(15 / 25.4)
        section.page_width = Inches(210 / 25.4)
        section.page_height = Inches(297 / 25.4)

    # Cấu hình Style mặc định
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # 2. PHẦN TIÊU NGỮ & QUỐC HIỆU (Bảng ẩn 2 cột chuẩn văn thư)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.8)
    table.columns[1].width = Inches(4.2)
    
    # Cột 1: Đơn vị công tác
    cell_left = table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.add_run("SỞ GIÁO DỤC VÀ ĐÀO TẠO\nTRƯỜNG THCS ...").bold = True
    p_left.runs[0].font.size = Pt(12)
    
    # Cột 2: Quốc hiệu - Tiêu ngữ
    cell_right = table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_r1 = p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    r_r1.bold = True
    r_r1.font.size = Pt(12)
    r_r2 = p_right.add_run("Độc lập - Tự do - Hạnh phúc\n_______________")
    r_r2.bold = True
    r_r2.font.size = Pt(13)

    # Khoảng trống xuống tên đề tài
    doc.add_paragraph().paragraph_format.space_before = Pt(24)

    # 3. TÊN ĐỀ TÀI SÁNG KIẾN KINH NGHIỆM
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(18)
    r_t1 = p_title.add_run("BÁO CÁO SÁNG KIẾN KINH NGHIỆM (BẢN ĐÃ HIỆU ĐÍNH)\n")
    r_t1.bold = True
    r_t1.font.size = Pt(14)
    
    clean_topic = topic.replace('"', '').replace('“', '').replace('”', '').replace('*', '').strip()
    r_t2 = p_title.add_run(f"“{clean_topic.upper()}”")
    r_t2.bold = True
    r_t2.font.size = Pt(15)

    # 4. QUÉT BIÊN DỊCH NỘI DUNG NÂNG CAO
    lines = refined_text.split('\n')
    for line in lines:
        cleaned_line = line.strip()
        if not cleaned_line or cleaned_line == "---":
            continue
            
        if cleaned_line.startswith('#'):
            cleaned_line = re.sub(r'^#+\s*', '', cleaned_line)
            is_heading = True
        else:
            is_heading = False
            
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        is_bullet = False
        if cleaned_line.startswith(('* ', '- ', '• ')):
            cleaned_line = re.sub(r'^[\*\-\•]\s*', '', cleaned_line)
            is_bullet = True
            p.paragraph_format.left_indent = Inches(0.4)
        
        is_admin_heading = cleaned_line.upper().startswith(("ĐẶT VẤN ĐỀ", "GIẢI PHÁP", "NỘI DUNG", "KẾT LUẬN", "KẾT QUẢ")) or re.match(r'^(\d+|\d+\.\d+|[I|V|X]+)\.\s*', cleaned_line)
        
        if is_heading or is_admin_heading:
            p.paragraph_format.space_before = Pt(12)
            if re.match(r'^\d+\.\d+', cleaned_line):
                p.paragraph_format.first_line_indent = Inches(0.5)
        elif not is_bullet:
            p.paragraph_format.first_line_indent = Inches(0.5)

        # Xử lý bóc tách in đậm cục bộ (Inline bold **) an toàn
        parts = re.split(r'(\*\*.*?\*\*)', cleaned_line)
        for part in parts:
            if not part:  # Bỏ qua nếu part trống để tránh lỗi UnboundLocalError
                continue
                
            if part.startswith('**') and part.endswith('**'):
                text_run = part.replace('**', '')
                run = p.add_run(text_run)
                run.bold = True
            else:
                run = p.add_run(part)
                if is_heading or is_admin_heading:
                    run.bold = True
            
            # Đảm bảo cấu hình font luôn đi liền sau khi 'run' chắc chắn tồn tại
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def show_docx_processor_module(api_key=None):
    st.subheader("📂 Trợ Lý AI Phân Tích & Thẩm Định Bản Nháp Word")
    st.info("Hỗ trợ đọc file .docx, tự động phát hiện lỗi chính tả, câu từ sư phạm và đánh giá cấu trúc theo quy chuẩn.")
    
    uploaded_file = st.file_uploader("Tải lên bản nháp Sáng kiến kinh nghiệm của thầy (.docx)", type=["docx"])
    
    if uploaded_file is not None:
        st.success("✅ Đã tải tệp lên thành công!")
        
        doc_title = st.text_input("Xác nhận hoặc nhập Tên Đề Tài để xuất file (.docx):", value="Ứng dụng công nghệ đổi mới phương pháp dạy học")
        
        if st.button("🧠 Kích hoạt AI quét và kiểm tra toàn diện"):
            with st.spinner("AI đang tiến hành thẩm định và hiệu đính văn bản song song..."):
                try:
                    # 1. Đọc nội dung file Word
                    doc = Document(uploaded_file)
                    full_text = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text)
                    
                    text_content = "\n".join(full_text)
                    preview_text = text_content[:8000]
                    
                    if len(text_content) == 0:
                        st.error("❌ File Word của thầy không có nội dung chữ.")
                        return

                    st.info(f"📋 Hệ thống đã ghi nhận đoạn văn bản dài {len(text_content)} ký tự.")
                    
                    # 2. Tạo Prompt tích hợp gộp luồng để tránh xung đột Streamlit
                    unified_prompt = f"""
                    Bạn là Chủ tịch Hội đồng Khoa học Giáo dục cấp huyện/tỉnh, chuyên gia phản biện đề tài nghiên cứu khoa học sư phạm ứng dụng và chuyên gia hiệu đính văn bản học thuật.
                    
                    Nhiệm vụ của bạn là đánh giá, phản biện và hiệu đính bản thảo sáng kiến kinh nghiệm hoặc đề tài nghiên cứu khoa học dưới góc nhìn của hội đồng xét duyệt.
                    
                    =========================
                    NỘI DUNG BẢN THẢO
                    =================
                    
                    {preview_text}
                    
                    =========================
                    YÊU CẦU ĐẦU RA
                    (BẮT BUỘC GIỮ NGUYÊN CÁC THẺ PHÂN TÁCH)
                    =======================================
                    
                    [PHAN_1_THAM_DINH]
                    
                    Hãy lập báo cáo thẩm định khoa học chi tiết theo các tiêu chí sau:
                    
                    1. 📝 ĐÁNH GIÁ VĂN PHONG HỌC THUẬT VÀ SƯ PHẠM
                    
                    * Mức độ chuẩn mực của ngôn ngữ hành chính giáo dục.
                    * Lỗi chính tả, ngữ pháp, diễn đạt.
                    * Các câu văn dài, khó hiểu hoặc mang tính khẩu ngữ.
                    * Những cụm từ sáo rỗng cần thay thế.
                    
                    2. 📐 ĐÁNH GIÁ CẤU TRÚC KHOA HỌC
                    
                    * Tính logic giữa các phần.
                    * Sự liên kết giữa thực trạng, giải pháp và kết quả.
                    * Mức độ đầy đủ của các nội dung.
                    * Các phần còn thiếu hoặc chưa đạt yêu cầu.
                    
                    3. 💡 ĐÁNH GIÁ TÍNH MỚI VÀ TÍNH SÁNG TẠO
                    
                    * Điểm mới của đề tài.
                    * Mức độ đổi mới phương pháp.
                    * Việc ứng dụng công nghệ số, AI hoặc chuyển đổi số.
                    * Khả năng nhân rộng và áp dụng thực tế.
                    
                    4. 📊 ĐÁNH GIÁ MINH CHỨNG VÀ THỰC NGHIỆM
                    
                    * Tính hợp lý của số liệu.
                    * Mức độ thuyết phục của minh chứng.
                    * Thiếu hụt bảng biểu, khảo sát, biểu đồ hoặc phụ lục.
                    * Đề xuất các minh chứng cần bổ sung.
                    
                    5. 🎯 ĐÁNH GIÁ THEO GÓC NHÌN HỘI ĐỒNG CHẤM
                       Chấm điểm theo thang 100:
                    
                    * Tính cấp thiết: .../10
                    * Cơ sở lý luận: .../10
                    * Tính mới: .../20
                    * Giải pháp thực hiện: .../25
                    * Hiệu quả thực tế: .../20
                    * Khả năng nhân rộng: .../10
                    * Hình thức trình bày: .../5
                    
                    Tổng điểm: .../100
                    
                    Xếp loại:
                    
                    * Xuất sắc
                    * Tốt
                    * Khá
                    * Đạt
                    * Chưa đạt
                    
                    6. 🛠 ĐỀ XUẤT CHỈNH SỬA
                    
                    * Những nội dung cần bổ sung.
                    * Các đoạn cần viết lại.
                    * Các số liệu cần minh chứng thêm.
                    * Các giải pháp cần làm rõ.
                    * Gợi ý nâng cao chất lượng để đạt cấp cao hơn.
                    
                    ---
                    
                    [PHAN_2_HIEU_DINH]
                    
                    Viết lại toàn bộ văn bản trên theo các yêu cầu:
                    
                    * Giữ nguyên nội dung cốt lõi.
                    * Sửa toàn bộ lỗi chính tả, ngữ pháp.
                    * Chuyển sang văn phong hành chính sư phạm học thuật.
                    * Tăng tính logic, tính khoa học và tính thuyết phục.
                    * Loại bỏ câu văn dài, lặp ý hoặc diễn đạt chưa chuẩn.
                    * Nếu nội dung còn thiếu dẫn chứng, được phép bổ sung các ví dụ hoặc số liệu giả định hợp lý.
                    * Không viết lời nhận xét, không giải thích, không thêm tiêu đề ngoài nội dung đã hiệu đính.
                    
                    ---
                    
                    [PHAN_3_GOI_Y_NANG_CAP]
                    
                    Đề xuất tối đa 10 giải pháp giúp đề tài nâng lên mức xét duyệt cao hơn như:
                    
                    * Bổ sung minh chứng.
                    * Ứng dụng AI hoặc chuyển đổi số.
                    * Thiết kế phiếu khảo sát.
                    * Bảng số liệu trước – sau.
                    * Biểu đồ thống kê.
                    * Phụ lục minh họa.
                    * Kế hoạch triển khai.
                    * Khả năng nhân rộng.
                    
                    Chỉ trình bày ngắn gọn theo dạng danh sách.
                    """

                    
                    # Gọi AI 1 lần duy nhất
                    ai_total_response = call_ai_stream(
                        prompt=unified_prompt,
                        system_instruction="Bạn là Giáo sư, Viện trưởng Viện Khoa học Giáo dục Việt Nam.",
                        api_key=api_key
                    )
                    
                    # 3. Kỹ thuật bóc tách chuỗi (String Splitting) để đưa vào giao diện
                    if "[PHAN_1_THAM_DINH]" in ai_total_response and "[PHAN_2_HIEU_DINH]" in ai_total_response:
                        parts = ai_total_response.split("[PHAN_2_HIEU_DINH]")
                        analysis_part = parts[0].replace("[PHAN_1_THAM_DINH]", "").strip()
                        refined_part = parts[1].strip()
                    else:
                        # Phương án dự phòng nếu AI quên ghi tag
                        analysis_part = ai_total_response
                        refined_part = ai_total_response

                    # 4. Hiển thị mượt mà lên giao diện từng phần tách biệt
                    st.markdown("### 📊 1. Kết quả Thẩm định từ Trợ lý AI:")
                    st.markdown(analysis_part)
                    
                    st.markdown("---")
                    st.markdown("### 📝 2. Bản thảo Sáng kiến sau khi AI tự động sửa đổi trực tiếp:")
                    st.markdown(refined_part)
                    
                    # Đóng gói xuất file Word
                    docx_bytes = export_refined_docx(doc_title, refined_part)
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    st.download_button(
                        label="📥 Tải Bản Sáng Kiến Đã Chỉnh Sửa Hoàn Thiện (.docx)",
                        data=docx_bytes,
                        file_name=f"SKKN_Da_Chinh_Sua_{doc_title.replace(' ', '_')[:25]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.success("✨ Đã đóng gói thành công file Word chuẩn thể thức!")
                    
                except Exception as e:
                    st.error(f"❌ Có lỗi xảy ra khi xử lý cấu trúc file Word: {str(e)}")
