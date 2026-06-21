import streamlit as st
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Import hàm gọi AI từ lõi hệ thống
try:
    from ai_engine import call_ai_stream
except ModuleNotFoundError:
    from edu_research_assistant.ai_engine import call_ai_stream

def export_to_docx_admin_format(topic, content_text):
    """
    Xuất nội dung văn bản sang file Word chuẩn định dạng hành chính Nghị định 30/2020/NĐ-CP
    Đã tích hợp bộ lọc bóc tách ký tự Markdown (** , ### , * ) tránh lỗi hiển thị chữ thô.
    """
    import re
    doc = Document()
    
    # 1. Cấu hình lề trang chuẩn (Top 20mm, Bottom 20mm, Left 30mm, Right 15mm)
    for section in doc.sections:
        section.top_margin = Inches(20 / 25.4)
        section.bottom_margin = Inches(20 / 25.4)
        section.left_margin = Inches(30 / 25.4)
        section.right_margin = Inches(15 / 25.4)
        section.page_width = Inches(210 / 25.4)
        section.page_height = Inches(297 / 25.4)

    # Cấu hình Style mặc định (Times New Roman, 14pt)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # --- PHẦN TIÊU NGỮ & QUỐC HIỆU (Bảng ẩn 2 cột) ---
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.8)
    table.columns[1].width = Inches(4.2)
    
    # Cột 1: Đơn vị
    cell_left = table.cell(0, 0)
    p_left1 = cell_left.paragraphs[0]
    p_left1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_l1 = p_left1.add_run("SỞ GIÁO DỤC VÀ ĐÀO TẠO\nTRƯỜNG THCS ...")
    run_l1.font.name = 'Times New Roman'
    run_l1.font.size = Pt(12)
    run_l1.bold = True
    
    p_left2 = cell_left.add_paragraph()
    p_left2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left2.paragraph_format.space_before = Pt(4)
    p_left2.add_run("-------").font.size = Pt(12)
    
    # Cột 2: Quốc hiệu Tiêu ngữ
    cell_right = table.cell(0, 1)
    p_right1 = cell_right.paragraphs[0]
    p_right1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_r1 = p_right1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    run_r1.bold = True
    run_r1.font.size = Pt(12)
    
    run_r2 = p_right1.add_run("Độc lập - Tự do - Hạnh phúc")
    run_r2.bold = True
    run_r2.font.size = Pt(13)
    
    p_right2 = cell_right.add_paragraph()
    p_right2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right2.paragraph_format.space_before = Pt(2)
    p_right2.add_run("_______________").bold = True

    # Khoảng cách xuống tên đề tài
    doc.add_paragraph().paragraph_format.space_before = Pt(24)

    # --- TÊN ĐỀ TÀI SÁNG KIẾN ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(18)
    r_title_lbl = p_title.add_run("BÁO CÁO SÁNG KIẾN KINH NGHIỆM\n")
    r_title_lbl.bold = True
    r_title_lbl.font.size = Pt(14)
    
    # Làm sạch tên đề tài nếu dính dấu ngoặc kép hoặc dấu sao của AI
    clean_topic = topic.replace('"', '').replace('“', '').replace('”', '').replace('*', '').strip()
    r_title_val = p_title.add_run(f"“{clean_topic.upper()}”")
    r_title_val.bold = True
    r_title_val.font.size = Pt(15)

    # --- BỘ LỌC VÀ BIÊN DỊCH VĂN BẢN CHUẨN ---
    lines = content_text.split('\n')
    for line in lines:
        cleaned_line = line.strip()
        if not cleaned_line or cleaned_line == "---":
            continue
            
        # 1. Khử toàn bộ các ký tự Tiêu đề Markdown (###, ##, #)
        if cleaned_line.startswith('#'):
            cleaned_line = re.sub(r'^#+\s*', '', cleaned_line)
            is_heading = True
        else:
            is_heading = False
            
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.3  # Giãn dòng 1.3 lines chuẩn
        p.paragraph_format.space_after = Pt(6)   # Cách đoạn dưới 6pt
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Căn đều hai bên
        
        # 2. Xử lý định dạng mục Đầu dòng / Danh sách (Bullet points hoặc dấu *)
        is_bullet = False
        if cleaned_line.startswith(('* ', '- ', '• ')):
            cleaned_line = re.sub(r'^[\*\-\•]\s*', '', cleaned_line)
            is_bullet = True
            p.paragraph_format.left_indent = Inches(0.4) # Thụt lề đoạn danh sách
        
        # Định nghĩa quy tắc in đậm cho tiêu đề lớn chuẩn hành chính
        is_admin_heading = cleaned_line.upper().startswith(("ĐẶT VẤN ĐỀ", "GIẢI PHÁP", "NỘI DUNG", "KẾT LUẬN", "KẾT QUẢ")) or re.match(r'^(\d+|\d+\.\d+|[I|V|X]+)\.\s*', cleaned_line)
        
        if is_heading or is_admin_heading:
            p.paragraph_format.space_before = Pt(12)
            # Tự động thụt lề 1.25cm nếu là tiêu đề tiểu mục nhỏ (ví dụ: 2.1. Lý do chọn đề tài)
            if re.match(r'^\d+\.\d+', cleaned_line):
                p.paragraph_format.first_line_indent = Inches(0.5)
        elif not is_bullet:
            # Thụt lề dòng đầu tiên 1.25cm đối với đoạn văn thường
            p.paragraph_format.first_line_indent = Inches(0.5)

        # 3. Sử dụng Regex bóc tách chuỗi **in đậm** ở giữa câu do AI sinh ra
        # Tách chuỗi theo cặp dấu **
        parts = re.split(r'(\*\*.*?\*\*)', cleaned_line)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Phần chữ nằm trong cặp dấu ** -> Bỏ dấu sao và đặt thuộc tính in đậm
                text_run = part.replace('**', '')
                run = p.add_run(text_run)
                run.bold = True
            else:
                # Phần chữ thường
                if part:
                    run = p.add_run(part)
                    if is_heading or is_admin_heading:
                        run.bold = True
            
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            
    # Ghi dữ liệu ra bộ nhớ đệm bytes
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def show_library_module(api_key=None):
    st.subheader("📚 Thư Viện Tra Cứu & Tự Động Sinh Sáng Kiến Mẫu")
    st.info("Thầy/Cô có thể tra cứu các mẫu có sẵn hoặc gõ một tên đề tài bất kỳ để AI tự động xây dựng một bản sáng kiến mẫu hoàn chỉnh.")

    # --- TÍNH NĂNG CAO CẤP: AI TỰ SINH SÁNG KIẾN MẪU THEO TÊN NGƯỜI DÙNG GÕ ---
    st.markdown("### 🤖 Tính năng: Tự động biên soạn Sáng kiến mẫu theo yêu cầu")
    
    user_topic = st.text_input(
        "✍️ Nhập tên Đề tài/Sáng kiến thầy cô muốn tạo mẫu:",
        placeholder="Ví dụ: Ứng dụng AI và chuyển đổi số nâng cao hiệu quả dạy học môn Toán lớp 8"
    )
    
    if st.button("🔥 Kích hoạt AI biên soạn Bản mẫu Sư phạm"):
        if not user_topic.strip():
            st.warning("⚠️ Vui lòng nhập tên đề tài trước khi bấm tạo mẫu.")
        else:
            with st.spinner("🧠 AI đang phân tích bối cảnh, tra cứu thuật ngữ và biên soạn bản mẫu..."):
                prompt = f"""
                Hãy đóng vai là Chuyên gia Khoa học Sư phạm cấp cao của Bộ Giáo dục và Đào tạo Việt Nam.
                Xây dựng một BẢN SÁNG KIẾN KINH NGHIỆM MẪU chi tiết cho đề tài: "{user_topic}".
                
                Yêu cầu cấu trúc nội dung sinh ra phải bao gồm đầy đủ các phần sau:
                1. TÊN ĐỀ TÀI (Chuẩn hóa lại nếu cần)
                2. ĐẶT VẤN ĐỀ
                   - Lý do chọn đề tài (Nêu bật tính cấp thiết và thực trạng khó khăn tại trường học).
                   - Mục đích nghiên cứu.
                3. GIẢI PHÁP THỰC HIỆN (Biên soạn chi tiết ít nhất 3 biện pháp cốt lõi, đột phá, có ứng dụng công nghệ hoặc đổi mới phương pháp).
                4. KẾT QUẢ ĐẠT ĐƯỢC (Đưa ra các số liệu định lượng, tỷ lệ % tăng trưởng giả định một cách khoa học).
                5. KẾT LUẬN & KIẾN NGHỊ (Bài học kinh nghiệm và đề xuất với nhà trường/Sở GD&ĐT).
                
                Văn phong yêu cầu: Nghiêm túc, mang tính học thuật ngành giáo dục Việt Nam, lập luận chặt chẽ, không sáo rỗng.
                """
                
                generated_sample = call_ai_stream(
                    prompt=prompt, 
                    system_instruction="Bạn là Giáo sư Viện Khoa học Giáo dục Việt Nam.", 
                    api_key=api_key
                )
                
                st.success("🎉 Đã biên soạn thành công bản mẫu sư phạm!")
                st.markdown("---")
                st.markdown(generated_sample)
                st.markdown("---")
                
                # --- TỰ ĐỘNG XỬ LÝ VÀ CHUYỂN ĐỔI SANG FILE WORD (.DOCX) CHUẨN ---
                try:
                    docx_bytes = export_to_docx_admin_format(user_topic, generated_sample)
                    
                    # Đổi nút bấm từ tải file .txt thành nút tải file .docx chuyên nghiệp
                    st.download_button(
                        label="📥 Tải Bản mẫu Word (.docx) Chuẩn Nghị Định 30",
                        data=docx_bytes,
                        file_name=f"Mau_SKKN_{user_topic.replace(' ', '_')[:30]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Lỗi khi đóng gói cấu trúc file Word hành chính: {e}")
                    # Nút tải file dự phòng dạng .txt nếu máy chủ bị nghẽn thư viện biên dịch
                    st.download_button(
                        label="📥 Tải Bản mẫu dự phòng (.txt)",
                        data=generated_sample,
                        file_name=f"Mau_SKKN_{user_topic.replace(' ', '_')[:30]}.txt",
                        mime="text/plain"
                    )
                
    st.markdown("<br><hr><br>", unsafe_allow_html=True)

    # --- THƯ VIỆN CÁC MẪU CỐ ĐỊNH CŨ ---
    st.markdown("### 📁 Danh mục các mẫu quy chuẩn có sẵn trong thư viện")
    search_query = st.text_input("🔍 Tìm kiếm nhanh trong kho mẫu sẵn có:", "")
    
    library_data = {
        "Toán học": [
            {
                "title": "Biện pháp ứng dụng công nghệ số trong dạy học Hình học lớp 8",
                "author": "Mẫu chuẩn Bộ GD&ĐT",
                "summary": "Giải pháp tập trung vào việc sử dụng các phần mềm hình học động (Geogebra) kết hợp mô hình lớp học đảo ngược...",
                "content": "Kính gửi: Quý Ban Giám hiệu, Hội đồng Khoa học Sư phạm, Phòng Giáo dục và Đào tạo, Sở Giáo dục và Đào tạo.

Tôi xin trình bày BẢN SÁNG KIẾN KINH NGHIỆM MẪU chi tiết với tư cách là Chuyên gia Khoa học Sư phạm cấp cao của Bộ Giáo dục và Đào tạo Việt Nam, về đề tài "Một số biện pháp khai thác và ứng dụng công nghệ số nhằm nâng cao hiệu quả dạy học và phát triển năng lực học sinh trong môn Hình học lớp 8".

BẢN SÁNG KIẾN KINH NGHIỆM MẪU
1. TÊN ĐỀ TÀI
Một số biện pháp khai thác và ứng dụng công nghệ số nhằm nâng cao hiệu quả dạy học và phát triển năng lực học sinh trong môn Hình học lớp 8.

2. ĐẶT VẤN ĐỀ
2.1. Lý do chọn đề tài

Chương trình Giáo dục phổ thông 2018 đang trong quá trình triển khai đã đặt ra yêu cầu đổi mới căn bản và toàn diện giáo dục, chuyển từ truyền thụ kiến thức sang phát triển phẩm chất và năng lực người học. Trong bối cảnh cuộc cách mạng công nghiệp 4.0 và xu thế chuyển đổi số quốc gia, việc tích hợp công nghệ số vào mọi mặt đời sống, đặc biệt là giáo dục, trở thành một nhiệm vụ cấp thiết và chiến lược.

Môn Hình học lớp 8, theo chương trình hiện hành, là giai đoạn chuyển tiếp quan trọng trong việc hình thành và phát triển tư duy hình học, khả năng hình dung không gian, lập luận logic và giải quyết vấn đề cho học sinh. Tuy nhiên, thực trạng dạy và học Hình học tại nhiều trường học vẫn còn đối mặt với những thách thức đáng kể:

Tính trừu tượng cao: Nhiều khái niệm, định lý hình học (như các loại tứ giác, đa giác, diện tích, thể tích, các phép biến hình) đòi hỏi khả năng hình dung, tưởng tượng tốt, điều mà không phải học sinh nào cũng có được một cách tự nhiên. Việc minh họa bằng bảng, phấn truyền thống thường tĩnh, một chiều, khó giúp học sinh hình dung các mối quan hệ động, các trường hợp tổng quát hay đặc biệt của hình.
Thiếu động lực học tập: Do tính trừu tượng và đôi khi khô khan của kiến thức, nhiều học sinh cảm thấy khó khăn, mất hứng thú, ngại học Hình học, dẫn đến kết quả học tập chưa cao và sự phát triển năng lực tư duy hình học bị hạn chế.
Hạn chế về phương pháp giảng dạy: Một bộ phận giáo viên vẫn quen với phương pháp truyền thụ một chiều, ít có cơ hội tiếp cận và ứng dụng các phương pháp dạy học tích cực có sự hỗ trợ của công nghệ số, hoặc chưa khai thác triệt để tiềm năng của các công cụ số sẵn có.
Thực trạng cơ sở vật chất: Mặc dù đã có sự đầu tư, nhưng việc ứng dụng công nghệ số một cách đồng bộ và hiệu quả trong dạy học Hình học còn gặp rào cản về hạ tầng công nghệ (máy tính, kết nối internet, phần mềm bản quyền) và kỹ năng khai thác của cả giáo viên lẫn học sinh.
Trong bối cảnh đó, việc nghiên cứu và ứng dụng công nghệ số trong dạy học Hình học lớp 8 không chỉ là một yêu cầu tất yếu để theo kịp xu thế thời đại mà còn là giải pháp hữu hiệu để khắc phục những hạn chế hiện có, tạo môi trường học tập sinh động, trực quan, khơi gợi hứng thú, phát huy tối đa năng lực tư duy và khả năng giải quyết vấn đề cho học sinh. Đề tài này được lựa chọn với kỳ vọng góp phần nâng cao chất lượng dạy học môn Hình học lớp 8, đồng thời là một minh chứng cụ thể cho định hướng chuyển đổi số trong giáo dục phổ thông.

2.2. Mục đích nghiên cứu

Đề tài nhằm đạt được các mục tiêu sau:

Nghiên cứu và đề xuất các biện pháp sư phạm cụ thể, khả thi trong việc khai thác và ứng dụng công nghệ số vào quá trình dạy học môn Hình học lớp 8 theo định hướng phát triển năng lực học sinh.
Thiết kế và thử nghiệm các hoạt động dạy học có ứng dụng công nghệ số nhằm nâng cao tính trực quan, sinh động, kích thích hứng thú học tập và sự chủ động, tích cực của học sinh.
Đánh giá hiệu quả của các biện pháp đã đề xuất thông qua việc cải thiện kết quả học tập, phát triển năng lực đặc thù môn Toán (tư duy hình học, lập luận, giải quyết vấn đề) và các năng lực chung (tin học, giao tiếp, hợp tác) của học sinh lớp 8.
Góp phần xây dựng một nguồn tài liệu tham khảo thiết thực, có giá trị ứng dụng cao cho giáo viên Hình học cấp trung học cơ sở trong quá trình đổi mới phương pháp giảng dạy.
3. GIẢI PHÁP THỰC HIỆN
Để đạt được các mục tiêu trên, tôi đã tập trung xây dựng và triển khai ba biện pháp cốt lõi, mang tính đột phá và ứng dụng công nghệ số mạnh mẽ như sau:

3.1. Biện pháp 1: Khai thác phần mềm hình học động (GeoGebra, Desmos) để trực quan hóa, mô phỏng và hỗ trợ quá trình khám phá kiến thức mới.

Cơ sở lý luận: Hình học động cung cấp môi trường tương tác, cho phép học sinh tự mình thay đổi các đối tượng hình học, quan sát sự biến đổi của chúng và rút ra các nhận xét, phỏng đoán. Điều này thúc đẩy tư duy khám phá, xây dựng kiến thức một cách chủ động thay vì thụ động tiếp thu.
Cách thức thực hiện:
Tập huấn cho học sinh và giáo viên: Tổ chức các buổi hướng dẫn cơ bản về cách sử dụng GeoGebra và Desmos để vẽ các hình cơ bản, đo đạc, thực hiện các phép biến hình.
Thiết kế các hoạt động khởi động và hình thành kiến thức mới:
Khám phá tính chất hình học: Ví dụ, khi dạy về các loại tứ giác (hình thang, hình bình hành, hình chữ nhật, hình thoi), giáo viên chuẩn bị sẵn các hình vẽ động. Học sinh có thể di chuyển các đỉnh, thay đổi kích thước để kiểm tra các tính chất về cạnh, góc, đường chéo, từ đó tự mình phát hiện ra các định lý, tính chất.
Trực quan hóa các khái niệm trừu tượng: Khi học về phép biến hình (phép đối xứng trục, đối xứng tâm, phép tịnh tiến), học sinh có thể trực tiếp kéo, xoay, tịnh tiến một hình để thấy được ảnh của nó, hiểu rõ bản chất của từng phép biến hình thay vì chỉ nhìn hình vẽ tĩnh.
Chứng minh định lý: Sử dụng GeoGebra để minh họa trực quan các bước chứng minh. Ví dụ, chứng minh định lý Pythagoras bằng cách ghép các hình vuông, hoặc chứng minh định lý Talet qua việc di chuyển các điểm trên đường thẳng song song.
Sử dụng trong củng cố và luyện tập: Giáo viên có thể tạo các bài tập tương tác trên GeoGebra, yêu cầu học sinh di chuyển điểm để đạt được một điều kiện nào đó, hoặc kiểm tra các mối quan hệ hình học.
Ưu điểm đột phá: Chuyển đổi hoàn toàn phương pháp giảng dạy từ truyền thụ sang kiến tạo. Học sinh không chỉ "thấy" mà còn "làm", "thử nghiệm" và "khám phá". Điều này đặc biệt hiệu quả trong việc phát triển năng lực tư duy hình học, khả năng suy luận và giải quyết vấn đề.
3.2. Biện pháp 2: Xây dựng và sử dụng hệ thống bài tập tương tác, kiểm tra đánh giá trực tuyến đa dạng trên các nền tảng công nghệ số.

Cơ sở lý luận: Việc ứng dụng công nghệ số trong đánh giá không chỉ giúp đa dạng hóa hình thức mà còn cung cấp phản hồi nhanh chóng, cá nhân hóa, giúp học sinh nhận diện lỗi sai và cải thiện kịp thời. Đồng thời, giáo viên có thể thu thập dữ liệu về quá trình học tập của học sinh một cách hiệu quả.
Cách thức thực hiện:
Sử dụng các nền tảng quiz tương tác: Kahoot, Quizizz, Google Forms, Microsoft Forms được sử dụng để tạo các bài kiểm tra trắc nghiệm khách quan, câu hỏi điền khuyết, kéo thả, ghép đôi, thường xuyên trong các hoạt động khởi động, củng cố bài học hoặc kiểm tra 15 phút.
Ví dụ: Thiết kế bài quiz nhanh về tính chất của hình bình hành sau khi học, hoặc nhận diện các dạng hình thông qua hình ảnh.
Tạo bài tập tự luận có hỗ trợ công cụ vẽ: Đối với các bài tập yêu cầu vẽ hình và chứng minh, có thể sử dụng các nền tảng LMS (ví dụ: Google Classroom, Microsoft Teams) để giao bài. Học sinh có thể vẽ hình bằng GeoGebra, chụp ảnh hoặc xuất file để nộp, sau đó giáo viên chấm bài và phản hồi trực tiếp trên hệ thống.
Phản hồi tự động và cá nhân hóa: Các nền tảng quiz cho phép thiết lập phản hồi tự động cho từng câu trả lời đúng/sai, giải thích chi tiết, hoặc gợi ý tài liệu học thêm. Điều này giúp học sinh học từ lỗi sai của mình một cách hiệu quả.
Tổng hợp dữ liệu và phân tích: Giáo viên có thể dễ dàng theo dõi tiến độ làm bài, kết quả của từng học sinh và của cả lớp, từ đó nhận diện được những phần kiến thức học sinh còn yếu, điều chỉnh kế hoạch giảng dạy phù hợp.
Ưu điểm đột phá: Tăng cường tính tương tác, hấp dẫn cho các hoạt động đánh giá. Giúp học sinh giảm áp lực thi cử, biến quá trình kiểm tra thành một phần của quá trình học tập. Cung cấp dữ liệu chi tiết cho giáo viên để cá nhân hóa việc hỗ trợ học sinh, nâng cao hiệu quả giáo dục.
3.3. Biện pháp 3: Tổ chức các dự án học tập, hoạt động STEM tích hợp công nghệ số nhằm gắn kết kiến thức Hình học với thực tiễn.

Cơ sở lý luận: Học tập dự án và STEM (Khoa học, Công nghệ, Kỹ thuật, Toán học) là phương pháp khuyến khích học sinh ứng dụng kiến thức liên môn để giải quyết các vấn đề thực tiễn, phát triển tư duy sáng tạo, kỹ năng làm việc nhóm và thuyết trình. Công nghệ số đóng vai trò là công cụ đắc lực để thực hiện và trình bày sản phẩm dự án.
Cách thức thực hiện:
Xây dựng chủ đề dự án liên môn:
Ví dụ 1 (Toán - Công nghệ - Mỹ thuật): "Thiết kế mô hình kiến trúc thân thiện môi trường trên nền tảng số". Học sinh sử dụng kiến thức về tỉ lệ, các hình phẳng, hình khối, diện tích, thể tích trong Hình học để thiết kế bản vẽ 2D (trên GeoGebra hoặc CAD đơn giản như SketchUp Free) và sau đó dựng mô hình 3D ảo. Sản phẩm có thể là một ngôi nhà, một công viên thu nhỏ.
Ví dụ 2 (Toán - Vật lí - Tin học): "Nghiên cứu ứng dụng của các phép biến hình trong thiết kế hoa văn, logo". Học sinh tìm hiểu các hoa văn truyền thống, logo có sử dụng phép đối xứng, tịnh tiến, quay. Sau đó, sử dụng GeoGebra hoặc phần mềm thiết kế đồ họa đơn giản để tạo ra các hoa văn, logo mới, đồng thời giải thích ý nghĩa hình học của chúng.
Sử dụng công nghệ số trong quá trình nghiên cứu và tạo sản phẩm: Học sinh tìm kiếm thông tin trên internet, sử dụng phần mềm để thiết kế, tính toán, mô phỏng, quay video, chụp ảnh, biên tập tài liệu.
Trình bày sản phẩm dự án bằng công nghệ số: Học sinh có thể tạo bài thuyết trình đa phương tiện (PowerPoint, Prezi), website đơn giản, video clip, hoặc trưng bày mô hình ảo trên các nền tảng trực tuyến.
Ưu điểm đột phá: Đưa kiến thức Hình học ra khỏi sách vở, gắn kết chặt chẽ với thế giới thực, giúp học sinh nhận thức được giá trị ứng dụng của môn học. Phát triển toàn diện các năng lực như giải quyết vấn đề, tư duy sáng tạo, phản biện, hợp tác, giao tiếp, khả năng thích ứng với công nghệ, và kỹ năng trình bày số hóa.
4. KẾT QUẢ ĐẠT ĐƯỢC
Sau một thời gian triển khai thử nghiệm các biện pháp trên tại các lớp 8 của một số trường học được chọn làm thí điểm (ví dụ: Trường THCS A, B), chúng tôi đã ghi nhận được những kết quả tích cực và mang tính định lượng rõ rệt:

4.1. Về chất lượng học tập môn Hình học:

Điểm trung bình môn Hình học: Đã có sự tăng trưởng đáng kể. Cụ thể, điểm trung bình của các lớp áp dụng sáng kiến tăng từ 6.5 lên 7.9 (tăng khoảng 21.5%), trong khi các lớp đối chứng chỉ tăng từ 6.4 lên 6.8 (tăng 6.25%).
Tỷ lệ học sinh đạt khá, giỏi: Tăng từ 48% lên 75%, tức là tăng 27 điểm phần trăm. Ngược lại, tỷ lệ này ở nhóm đối chứng chỉ tăng từ 46% lên 52%.
Tỷ lệ học sinh yếu, kém: Giảm mạnh từ 12% xuống dưới 3%. Trong khi đó, ở nhóm đối chứng, tỷ lệ này giảm từ 10% xuống 7%.
Khả năng giải quyết các bài toán hình học phức tạp: Ổn định và tăng lên rõ rệt, đặc biệt là các bài toán yêu cầu tư duy không gian và ứng dụng thực tiễn. Khoảng 70% học sinh thể hiện sự tự tin và khả năng vận dụng tốt kiến thức.
4.2. Về hứng thú và thái độ học tập:

Khảo sát định lượng về mức độ hứng thú: Hơn 88% học sinh tham gia các hoạt động ứng dụng công nghệ số thể hiện sự hứng thú "rất cao" hoặc "cao" đối với môn Hình học, tăng mạnh so với con số 55% trước khi áp dụng.
Mức độ tham gia tích cực: Tần suất học sinh phát biểu, đặt câu hỏi, thảo luận nhóm tăng khoảng 30% trong các tiết học có ứng dụng công nghệ số.
Thái độ chủ động, sáng tạo: Học sinh tích cực hơn trong việc tự tìm tòi, khám phá kiến thức và chủ động đề xuất các ý tưởng trong các dự án học tập.
4.3. Về sự phát triển năng lực của học sinh:

Năng lực tư duy hình học và lập luận: 78% học sinh được đánh giá có sự tiến bộ rõ rệt trong việc hình dung, phân tích các đối tượng hình học và đưa ra các lập luận logic.
Năng lực sử dụng công nghệ thông tin: Hơn 95% học sinh thành thạo các thao tác cơ bản trên phần mềm GeoGebra, Desmos và các nền tảng kiểm tra trực tuyến, biết cách tìm kiếm thông tin, sử dụng công cụ trình bày sản phẩm số.
Năng lực hợp tác, giao tiếp và giải quyết vấn đề: 85% học sinh thể hiện khả năng làm việc nhóm hiệu quả, biết cách phân công nhiệm vụ, trao đổi ý kiến và cùng nhau tìm ra giải pháp cho các vấn đề thực tiễn trong các dự án học tập.
4.4. Về sự chuyển biến trong phương pháp giảng dạy của giáo viên:

100% giáo viên tham gia thử nghiệm đều nhận thấy hiệu quả tích cực của việc ứng dụng công nghệ số và bày tỏ mong muốn tiếp tục áp dụng, phát triển các biện pháp này trong giảng dạy.
Giáo viên trở nên chủ động hơn trong việc tìm tòi, học hỏi các công cụ công nghệ mới và linh hoạt hơn trong việc thiết kế các hoạt động học tập.
Những kết quả trên cho thấy, việc ứng dụng công nghệ số trong dạy học Hình học lớp 8 không chỉ là một giải pháp mà còn là một xu thế tất yếu, mang lại hiệu quả vượt trội trong việc nâng cao chất lượng giáo dục và phát triển toàn diện năng lực học sinh theo định hướng của Chương trình GDPT 2018.

5. KẾT LUẬN & KIẾN NGHỊ
5.1. Kết luận (Bài học kinh nghiệm)

Qua quá trình nghiên cứu, xây dựng và triển khai các biện pháp ứng dụng công nghệ số trong dạy học Hình học lớp 8, chúng tôi rút ra một số bài học kinh nghiệm quan trọng sau:

Tính cấp thiết và hiệu quả vượt trội: Việc ứng dụng công nghệ số không còn là lựa chọn mà là yêu cầu bắt buộc để nâng cao hiệu quả dạy học, đặc biệt với các môn học có tính trực quan và trừu tượng cao như Hình học. Công nghệ số đã biến các khái niệm khô khan thành những trải nghiệm học tập sinh động, giúp học sinh tiếp cận kiến thức một cách sâu sắc và bền vững hơn.
Vai trò trung tâm của giáo viên: Công nghệ chỉ là công cụ. Thành công của việc ứng dụng phụ thuộc rất lớn vào sự chủ động, sáng tạo, không ngừng học hỏi và đổi mới của giáo viên. Giáo viên cần phải là người thiết kế, định hướng và dẫn dắt học sinh khám phá, khai thác tối đa tiềm năng của công nghệ.
Phát huy tính chủ động của học sinh: Khi được trao quyền và công cụ, học sinh sẽ phát huy tối đa tính chủ động, tò mò, khám phá và phát triển các năng lực cốt lõi. Môi trường học tập số tạo điều kiện cho học sinh tự học, tự điều chỉnh và hợp tác hiệu quả.
Yêu cầu về sự đồng bộ: Để các biện pháp ứng dụng công nghệ số phát huy hiệu quả tối đa, cần có sự đầu tư đồng bộ về cơ sở vật chất (thiết bị, đường truyền internet), phần mềm, cũng như các chương trình tập huấn, bồi dưỡng thường xuyên cho đội ngũ giáo viên.
Không thay thế mà hỗ trợ: Công nghệ số không thể và không nên thay thế hoàn toàn vai trò của giáo viên hay các phương pháp truyền thống. Thay vào đó, nó là công cụ hỗ trợ đắc lực, mở rộng không gian và khả năng sáng tạo trong dạy và học.
5.2. Kiến nghị

Để nhân rộng và phát huy hiệu quả của sáng kiến này, tôi xin đưa ra một số kiến nghị cụ thể:

Với Nhà trường (Ban Giám hiệu):

Tiếp tục đầu tư cơ sở vật chất: Trang bị đầy đủ các thiết bị công nghệ cần thiết (máy tính, máy chiếu đa năng, màn hình tương tác, kết nối internet tốc độ cao) cho các phòng học và phòng bộ môn.
Tổ chức tập huấn thường xuyên: Định kỳ tổ chức các buổi tập huấn chuyên sâu về khai thác các phần mềm hình học động (GeoGebra, Desmos) và các nền tảng kiểm tra, đánh giá trực tuyến cho giáo viên.
Khuyến khích và động viên: Tạo môi trường khuyến khích giáo viên chủ động nghiên cứu, thử nghiệm các biện pháp ứng dụng công nghệ số trong dạy học thông qua các hình thức khen thưởng, chia sẻ kinh nghiệm.
Xây dựng kho học liệu số: Tổ chức xây dựng một kho học liệu số dùng chung của nhà trường bao gồm các bài giảng điện tử, video hướng dẫn, bài tập tương tác được thiết kế bởi giáo viên trong trường.
Với Phòng Giáo dục và Đào tạo/Sở Giáo dục và Đào tạo:

Xây dựng lộ trình chuyển đổi số: Ban hành các văn bản hướng dẫn, lộ trình cụ thể về việc đẩy mạnh ứng dụng công nghệ số trong dạy và học các môn học, đặc biệt là các môn khoa học tự nhiên.
Tổ chức các chuyên đề, hội thảo: Định kỳ tổ chức các chuyên đề, hội thảo cấp huyện/tỉnh về ứng dụng công nghệ số trong dạy học Hình học và các môn học khác để giáo viên có cơ hội học hỏi, trao đổi kinh nghiệm.
Thẩm định và nhân rộng: Đánh giá, thẩm định các sáng kiến kinh nghiệm hiệu quả về ứng dụng công nghệ số để có kế hoạch nhân rộng trên địa bàn toàn tỉnh/thành phố.
Hỗ trợ hạ tầng công nghệ: Đề xuất chính sách hỗ trợ các trường vùng sâu, vùng xa, vùng khó khăn trong việc trang bị hạ tầng công nghệ thông tin phục vụ chuyển đổi số giáo dục.
Phát triển nền tảng học liệu chung: Đầu tư phát triển các nền tảng học liệu số chung cấp tỉnh, tích hợp các công cụ hỗ trợ dạy học và kiểm tra, đánh giá trực tuyến.
Tôi tin rằng, với sự quan tâm, đầu tư và định hướng đúng đắn từ các cấp quản lý giáo dục, cùng với sự nỗ lực, sáng tạo của đội ngũ giáo viên, việc ứng dụng công nghệ số trong dạy học Hình học lớp 8 nói riêng và các môn học khác nói chung sẽ ngày càng phát huy hiệu quả, góp phần hiện thực hóa mục tiêu đổi mới căn bản, toàn diện giáo dục Việt Nam.

Trân trọng./.

Chuyên gia Khoa học Sư phạm cấp cao Bộ Giáo dục và Đào tạo Việt Nam (Ký tên)"
            }
        ]
    }
    
    for cat, topics in library_data.items():
        filtered_topics = [t for t in topics if search_query.lower() in t["title"].lower() or search_query.lower() in cat.lower()]
        if filtered_topics:
            with st.expander(f"📁 Danh mục: Môn {cat} ({len(filtered_topics)} tài liệu)"):
                for t in filtered_topics:
                    st.markdown(f"### 📄 {t['title']}")
                    st.write(f"ℹ️ **Tóm tắt:** {t['summary']}")
                    if st.checkbox(f"👁️ Xem nội dung mẫu", key=t['title']):
                        st.text_area("Nội dung:", value=t['content'], height=150)
